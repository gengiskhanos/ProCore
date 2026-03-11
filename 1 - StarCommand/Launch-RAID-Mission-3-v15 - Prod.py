import configparser
import io
import json
import os
import re
import sys
import time
import tkinter as tk
import xml.etree.ElementTree as ET
import zipfile
from tkinter import messagebox, simpledialog

import oletools
import requests
import sseclient
from docx import Document
from docx.shared import Pt
from openpyxl import load_workbook

# ==========================================================
# CHARGEMENT DE LA CONFIGURATION (config.ini)
# ==========================================================
config = configparser.ConfigParser()

# Gestion du chemin pour compatibilité EXE (PyInstaller)
if getattr(sys, "frozen", False):
    app_path = os.path.dirname(sys.executable)
else:
    app_path = os.path.dirname(__file__)

config_path = os.path.join(app_path, "config.ini")

if not os.path.exists(config_path):
    # Création d'un template si manquant
    config["EDF_API"] = {"api_key": "CLE_ICI", "project_id": "ID_ICI"}
    config["PATHS"] = {"base_path": "CHEMIN_ICI"}
    config["SETTINGS"] = {"max_retry": "4", "wait_between_retry": "90"}
    with open(config_path, "w") as f:
        config.write(f)
    print(f"Fichier config.ini créé. Veuillez le remplir et relancer.")
    sys.exit()

config.read(config_path, encoding="utf-8")

API_KEY = config["EDF_API"]["api_key"]
PROJECT_ID = config["EDF_API"]["project_id"]
# --- MODIFICATION ICI ---
# On récupère la valeur du dossier (ex: "StarCommand-Docs-A-Generer")
relative_base = config["PATHS"]["base_path"]

# On le rend absolu en le joignant au chemin de l'application
# Si l'utilisateur met un chemin absolu (commençant par C:), os.path.join l'utilisera intelligemment.
BASE_PATH = os.path.abspath(os.path.join(app_path, relative_base))
# -------------------------

MAX_RETRY = int(config["SETTINGS"]["max_retry"])
WAIT_BETWEEN_RETRY = int(config["SETTINGS"]["wait_between_retry"])

print(f"📍 Chemin racine utilisé : {BASE_PATH}")


# ==========================================================
# INTERFACE DE SAISIE "UN CLIC"
# ==========================================================
def obtenir_num_dossier():
    root = tk.Tk()
    root.withdraw()
    num = simpledialog.askstring(
        "StarDoc ✨", "Veuillez saisir le numéro du dossier (ex: 378) :"
    )
    if not num:
        sys.exit()
    return num


NUM_DOSSIER = obtenir_num_dossier()
DOSSIER_DOCUMENTS = os.path.join(BASE_PATH, NUM_DOSSIER)

if not os.path.exists(DOSSIER_DOCUMENTS):
    messagebox.showerror("Erreur", f"Dossier introuvable :\n{DOSSIER_DOCUMENTS}")
    sys.exit()

# Paramètres API restants
URL_DOCUMENT = "https://api.iag.edf.fr/v2/workspaces/HcA-puQ/webhooks/document"
URL_QUERY = "https://api.iag.edf.fr/v2/workspaces/HcA-puQ/webhooks/query"
URL_DELETE = "https://api.iag.edf.fr/v2/workspaces/HcA-puQ/webhooks/document"
HEADERS = {"knowledge-project-apikey": API_KEY}
NOM_FICHIER_SORTIE = f"STARDOC-{NUM_DOSSIER}.docx"


# ==========================================================
# UTILITIES (retry)
# ==========================================================
def requete_avec_retry(method, url, **kwargs):
    for tentative in range(1, MAX_RETRY + 1):
        try:
            response = requests.request(method, url, timeout=180, **kwargs)
            if response.status_code in [200, 201]:
                return response
            else:
                print(
                    f"⚠ Tentative {tentative}/{MAX_RETRY} - Code {response.status_code}"
                )
                if response.status_code == 504:
                    print(f"⏳ Attente {WAIT_BETWEEN_RETRY}s...")
                    time.sleep(WAIT_BETWEEN_RETRY)
        except requests.exceptions.Timeout:
            print(f"⏳ Timeout - tentative {tentative}")
            time.sleep(WAIT_BETWEEN_RETRY)
        except Exception as e:
            print(f"❌ Erreur : {e}")
            time.sleep(WAIT_BETWEEN_RETRY)
    return None


# ==========================================================
# FONCTIONS DE PARSING DE FICHIERS (conversion pré-upload)
# ==========================================================
def parse_uipath_xaml(file_path):
    try:
        tree = ET.parse(file_path)
        root = tree.getroot()
        activities, variables = [], []
        for elem in root.iter():
            tag = elem.tag.split("}")[-1]
            if "DisplayName" in elem.attrib:
                activities.append({"activity": tag, "name": elem.attrib["DisplayName"]})
            if "x:Name" in elem.attrib or "Name" in elem.attrib:
                variables.append(
                    {
                        "name": elem.attrib.get("x:Name", elem.attrib.get("Name")),
                        "type": tag,
                    }
                )
        return {"type": "UiPath", "activities": activities, "variables": variables}
    except Exception as e:
        return {"error": str(e)}


def parse_power_automate_json(file_path):
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            data = json.load(f)
        actions = data.get("definition", {}).get("actions", {})
        steps = [
            {"name": k, "type": v.get("type"), "inputs": v.get("inputs", {})}
            for k, v in actions.items()
        ]
        return {"type": "Power Automate", "steps": steps}
    except Exception as e:
        return {"error": str(e)}


def parse_powerbi_pbix(file_path):
    """Extraction profonde des métadonnées Power BI pour le RAG"""
    try:
        with zipfile.ZipFile(file_path, "r") as z:
            info = {
                "type": "Power BI (Deep Parse)",
                "tables_found": [],
                "measures_hint": [],
            }
            # Analyse du DataModel (extraction des chaînes de texte)
            if "DataModel" in z.namelist():
                content = z.read("DataModel")
                # On cherche les motifs de texte (tables/colonnes) dans le binaire
                strings = re.findall(rb"[A-Z][a-zA-Z0-9_]{2,30}", content)
                unique_strings = sorted(
                    list(set([s.decode("utf-8", errors="ignore") for s in strings]))
                )
                info["data_model_vocabulary"] = unique_strings[:200]  # Top 200 termes

            # Analyse du Layout (visuels)
            if "Report/Layout" in z.namelist():
                layout = z.read("Report/Layout").decode("utf-16", errors="ignore")
                visuals = re.findall(r'"name":"([^"]+)"', layout)
                info["visual_elements"] = list(set(visuals))[:50]

            return info
    except Exception as e:
        return {"error": f"PBIX Parse Error: {str(e)}"}


def extract_vba_from_xlsm(file_path):
    """
    Extraction des macros VBA à partir d'un .xlsm.
    - Si oletools est installé, on extrait le code des modules.
    - Sinon, on indique seulement la présence de macros.
    """
    try:
        with zipfile.ZipFile(file_path, "r") as z:
            namelist = z.namelist()
            if "xl/vbaProject.bin" not in namelist:
                return {"has_vba": False, "modules": []}

            vba_data = z.read("xl/vbaProject.bin")

            try:
                # Tentative d'utilisation de oletools si disponible
                from oletools.olevba import VBA_Parser

                vba = VBA_Parser(filename="", data=vba_data)
                modules = []
                for (
                    filename,
                    stream_path,
                    vba_filename,
                    vba_code,
                ) in vba.extract_macros():
                    modules.append(
                        {
                            "module": vba_filename or filename,
                            "stream_path": " / ".join(stream_path)
                            if stream_path
                            else None,
                            "code": vba_code,
                        }
                    )
                vba.close()
                return {"has_vba": True, "modules": modules}
            except ImportError:
                # oletools non disponible : on indique juste la présence de macros
                return {
                    "has_vba": True,
                    "modules": [],
                    "warning": "oletools non installé, impossible d’extraire le code VBA",
                }
            except Exception as e:
                return {
                    "has_vba": True,
                    "modules": [],
                    "error": f"Erreur lors de l’extraction VBA : {str(e)}",
                }
    except Exception as e:
        return {
            "has_vba": False,
            "modules": [],
            "error": f"Erreur d’accès au vbaProject.bin : {str(e)}",
        }


def extract_excel_logic_universal(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    logic_report = {
        "filename": os.path.basename(file_path),
        "type": f"Excel ({ext})",
        "macros": extract_vba_from_xlsm(file_path) if ext == ".xlsm" else None,
        "potential_source_paths": [],
        "power_query_m": [],
        "ms_queries": [],
        "sheet_structures": {},
    }

    try:
        # --- 1. EXTRACTION BINAIRE (Power Query & Connections) ---
        with zipfile.ZipFile(file_path, "r") as z:
            # A. Récupération du Power Query (Code M)
            # On cherche dans les items customXml le flux DataMashup
            for item in [f for f in z.namelist() if "customXml/itemData" in f]:
                data = z.read(item)
                if b"PK\x03\x04" in data:  # Signature d'un ZIP imbriqué
                    start = data.find(b"PK\x03\x04")
                    try:
                        with zipfile.ZipFile(io.BytesIO(data[start:])) as sub_z:
                            if "Formulas/Section1.m" in sub_z.namelist():
                                m_code = sub_z.read("Formulas/Section1.m").decode(
                                    "utf-16-le", errors="ignore"
                                )
                                # On nettoie pour ne garder que les blocs let...in
                                logic_report["power_query_m"].extend(
                                    re.findall(
                                        r"let.*?in.*?(?=\r|\n|$)",
                                        m_code,
                                        re.DOTALL | re.IGNORECASE,
                                    )
                                )
                    except:
                        pass

            # B. Récupération des commandes SQL
            if "xl/connections.xml" in z.namelist():
                with z.open("xl/connections.xml") as f:
                    tree = ET.parse(f)
                    for conn in tree.findall(".//{*}connection"):
                        db_pr = conn.find(".//{*}dbPr")
                        if db_pr is not None:
                            logic_report["ms_queries"].append(
                                {"name": conn.get("name"), "sql": db_pr.get("command")}
                            )

        # --- 2. SCAN DYNAMIQUE DES FEUILLES (Détection intelligente) ---
        wb = load_workbook(file_path, data_only=False, read_only=True)
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]

            # 1. Trouver la première ligne non vide (Header Row)
            header_row_idx = 0
            headers = []
            for row in ws.iter_rows(min_row=1, max_row=10):
                row_values = [str(cell.value) for cell in row if cell.value is not None]
                if len(row_values) > 1:
                    headers = row_values
                    header_row_idx = row[0].row
                    break

            if not headers:
                continue

            # 2. Extraire les formules de la ligne juste en dessous (Data Row)
            formulas = {}
            data_row_idx = header_row_idx + 1

            # --- CORRECTION ICI : Sécurisation de l'accès à la ligne ---
            rows_generator = list(
                ws.iter_rows(min_row=data_row_idx, max_row=data_row_idx)
            )

            if rows_generator:  # On vérifie que la liste n'est pas vide
                data_row = rows_generator[0]

                for cell in data_row:
                    val = cell.value
                    if val and isinstance(val, str):
                        col_idx = cell.column - 1
                        # Sécurité sur l'index des headers
                        header_name = (
                            headers[col_idx]
                            if col_idx < len(headers)
                            else f"Col_{cell.column}"
                        )

                        # Détection Formule
                        if val.startswith("="):
                            formulas[header_name] = val

                        # Détection Chemin Réseau
                        elif "\\\\" in val or (":" in val and "\\" in val):
                            if val not in logic_report["potential_source_paths"]:
                                logic_report["potential_source_paths"].append(val)

            if headers or formulas:
                logic_report["sheet_structures"][sheet_name] = {
                    "header_detected_at_line": header_row_idx,
                    "columns": headers,
                    "formulas_samples": formulas,
                }
    except Exception as e:
        logic_report["error"] = f"Erreur lors de l'analyse : {str(e)}"

    return logic_report


def extract_powerquery_from_excel(file_path):
    try:
        queries = []
        with zipfile.ZipFile(file_path, "r") as z:
            for name in z.namelist():
                if name.startswith("customXml/") and name.endswith(".xml"):
                    with z.open(name) as f:
                        try:
                            tree = ET.parse(f)
                            root = tree.getroot()
                            for elem in root.iter():
                                if elem.text and (
                                    "let" in elem.text
                                    or "in" in elem.text
                                    or "#" in elem.text
                                ):
                                    queries.append(elem.text.strip())
                        except ET.ParseError:
                            continue
        return queries
    except Exception as e:
        return [f"Error extracting Power Query: {str(e)}"]


def parse_powerquery_m(file_path):
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            lines = f.readlines()
            steps = [line.strip() for line in lines if line.strip()]
        return {"type": "Power Query", "steps": steps}
    except Exception as e:
        return {"error": str(e)}


# ==========================================================
# CONVERSION AUTOMATIQUE AVANT UPLOAD
# ==========================================================
def convertir_si_necessaire(chemin_fichier):
    ext = os.path.splitext(chemin_fichier)[1].lower()

    # Extensions autorisées à être converties
    mapping = {
        ".xaml": parse_uipath_xaml,
        ".pbix": parse_powerbi_pbix,
        ".m": parse_powerquery_m,
        ".xlsm": extract_excel_logic_universal,
        ".xlsx": extract_excel_logic_universal,
    }

    if ext not in mapping:
        return chemin_fichier  # Pas de conversion

    print(f"🔄 Conversion : {os.path.basename(chemin_fichier)}")
    parser = mapping[ext]
    result = parser(chemin_fichier)

    chemin_json = chemin_fichier + "_parsed.json"
    with open(chemin_json, "w", encoding="utf-8") as f:
        json.dump(result, f, indent=2, ensure_ascii=False)

    return chemin_json


# ==========================================================
# MOTEUR MARKDOWN → WORD
# ==========================================================
def clean_html_noise(text):
    if not text:
        return ""
    text = re.sub(r"<pr-canvas[^>]*>", "", text)
    text = re.sub(r"</pr-canvas>", "", text)
    text = re.sub(r"<[^>]+>", "", text)
    text = text.replace("**", "")
    return text


def convert_markdown_table(table_text):
    rows = [r.strip() for r in table_text.strip().split("\n") if "|" in r]
    rows = [r.strip("|") for r in rows]
    rows = [r.split("|") for r in rows]

    if len(rows) < 2:
        return [], []

    header = rows[0]
    data = [r for r in rows[2:]]
    return header, data


def add_markdown_table_to_doc(doc, header, data):
    if not header and not data:
        return

    # Création du tableau dans Word
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"

    # 1. Remplissage de l'entête
    hdr_cells = table.rows[0].cells
    for i, text in enumerate(header):
        # On nettoie les étoiles de l'entête
        hdr_cells[i].text = text.replace("**", "").replace("`", "").strip()

    # 2. Remplissage des lignes de données
    for row_data in data:
        row_cells = table.add_row().cells
        for i, text in enumerate(row_data):
            # On nettoie les étoiles ET les backticks (`) pour les variables techniques
            clean_text = text.replace("**", "").replace("`", "").strip()
            row_cells[i].text = clean_text

    doc.add_paragraph()  # Petit espace après le tableau


def ajouter_texte_markdown(doc, texte):
    # 1. NETTOYAGE DES ESPACES ET CARACTERES SPECIAUX
    texte = texte.replace("\xa0", " ").replace("\u200b", "").strip()
    lignes = texte.split("\n")
    buffer_table = []

    # Fonction locale pour supprimer toutes les étoiles résiduelles
    def clean_stars(t):
        return t.replace("**", "").replace("*", "").strip()

    for ligne in lignes:
        ligne_raw = ligne.strip()
        if not ligne_raw:
            continue

        # --- GESTION DES TABLEAUX ---
        if "|" in ligne_raw and not ligne_raw.startswith("#"):
            buffer_table.append(ligne_raw)
            continue
        else:
            if buffer_table:
                header, data = convert_markdown_table("\n".join(buffer_table))
                add_markdown_table_to_doc(doc, header, data)
                buffer_table = []

        # --- GESTION DES TITRES # ---
        if ligne_raw.startswith("#"):
            titre = clean_stars(ligne_raw.lstrip("#"))
            doc.add_heading(titre, level=2)
            continue

        # --- GESTION DES TITRES DE SECTIONS (- **Titre**) ---
        # On cherche le format strict : Tiret + Texte entre étoiles
        match_titre_section = re.match(r"^\s*[-*]\s+\*\*(.*?)\*\*\s*$", ligne_raw)

        if match_titre_section:
            clean_title = clean_stars(match_titre_section.group(1))
            doc.add_paragraph()
            p = doc.add_paragraph()
            run = p.add_run(clean_title)
            run.bold = True
            run.font.size = Pt(13)
            doc.add_paragraph()
            continue

        # --- GESTION DES PUCES CLASSIQUES ---
        if ligne_raw.startswith("- ") or ligne_raw.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            # On nettoie la ligne de son tiret et de TOUTES ses étoiles
            texte_puce = clean_stars(ligne_raw.lstrip("-* "))
            p.add_run(texte_puce)
            continue

        # --- PARAGRAPHE NORMAL ---
        # On nettoie toutes les étoiles pour le corps du texte
        p = doc.add_paragraph()
        p.add_run(clean_stars(ligne_raw))

    # Vidage du buffer tableau si la réponse finit par un tableau
    if buffer_table:
        header, data = convert_markdown_table("\n".join(buffer_table))
        add_markdown_table_to_doc(doc, header, data)


# ==========================================================
# UPLOAD DES DOCUMENTS EDF (RAG)
# ==========================================================
def uploader_les_documents(dossier_path):
    print("--- Étape 1 : Envoi des fichiers (avec conversion automatique) ---")
    noms_documents_uploades = []  # Contiendra les NOMS (pour les prompts)
    ids_documents_rag = []  # Contiendra les IDs RAG (pour la suppression)

    for nom_fichier in os.listdir(dossier_path):
        chemin_complet = os.path.join(dossier_path, nom_fichier)
        if not os.path.isfile(chemin_complet):
            continue

        chemin_upload = convertir_si_necessaire(chemin_complet)
        nom_upload = os.path.basename(chemin_upload)

        print(f"Upload du fichier : {nom_upload}")
        with open(chemin_upload, "rb") as f:
            files = {"file": (nom_upload, f)}
            data = {"projectId": PROJECT_ID, "name": nom_upload}
            response = requete_avec_retry(
                "POST", URL_DOCUMENT, headers=HEADERS, data=data, files=files
            )

        doc_id = None
        if response:
            try:
                resp_json = response.json()
                doc_id = resp_json.get("_id")
                # Vérification au cas où l'API retourne une structure de recherche
                if (
                    not doc_id
                    and "search" in resp_json
                    and "results" in resp_json["search"]
                ):
                    results = resp_json["search"]["results"]
                    if results:
                        doc_id = results[0].get("_id")
            except Exception as e:
                print(f"❌ Erreur lors du parsing JSON : {e}")

            if doc_id:
                # Stockage des deux informations distinctement
                noms_documents_uploades.append(nom_upload)  # <-- Ajoute le nom
                ids_documents_rag.append(doc_id)  # <-- Ajoute l'ID RAG
                print(f"✅ Upload en RAG : {nom_upload} (ID: {doc_id})")
            else:
                print(
                    f"⚠ Fichier {nom_upload} non trouvé ou ID manquant dans la réponse API."
                )

    print("--- Upload terminé ---\n")
    return noms_documents_uploades, ids_documents_rag  # <-- Retourne les deux listes


# ==========================================================
# INTERROGER L’AGENT EDF
# ==========================================================
def interroger_agent_sse(question):
    print(f"➡ Question (Mode SSE) : {question}")

    body = {
        "text": question,
        "projectId": PROJECT_ID,
        "userId": "script_python_doc_gen",
        "sse": True,  # <--- Activation du Server-Side Events
        "projectConfigOverride": {"ai": {"history": True}},
    }

    # On utilise stream=True dans requests pour garder la connexion ouverte
    response = requests.post(URL_QUERY, headers=HEADERS, json=body, stream=True)

    if response.status_code != 200:
        print(f"Erreur : {response.status_code}")
        return ""

    client = sseclient.SSEClient(response)
    reponse_complete = ""

    # On boucle sur les événements envoyés par l'API
    for event in client.events():
        # L'API envoie chaque morceau (chunk) de texte [cite: 19]
        # Le format final est généralement un JSON dans event.data
        if event.data:
            try:
                data_json = json.loads(event.data)

                # Selon la doc, chaque message contient l'id et la réponse cumulée [cite: 19, 20]
                if "answer" in data_json:
                    reponse_complete = data_json["answer"]
                    # Optionnel : afficher l'avancement en temps réel
                    # print(reponse_complete, end="\r")

                # L'API envoie le JSON complet final comme dernier message [cite: 20]
                if "messageId" in data_json and not data_json.get("error"):
                    # On peut s'arrêter quand on a reçu l'objet final complet
                    pass
            except Exception:
                # Gestion des messages non-JSON ou fins de flux
                continue

    return reponse_complete


# ==========================================================
# SUPPRIMER DOCUMENTS DU RAG
# ==========================================================
def supprimer_documents(ids):
    print("--- Suppression des documents du RAG ---")
    for doc_id in ids:
        url = f"{URL_DELETE}?projectId={PROJECT_ID}&id={doc_id}"
        response = requete_avec_retry("DELETE", url, headers=HEADERS)
        if response:
            print(f"🗑 Supprimé : {doc_id}")
        else:
            print(f"⚠ Échec suppression : {doc_id}")
    print("--- Suppression terminée ---\n")


# ==========================================================
# SCRIPT PRINCIPAL
# ==========================================================

# 1. RÉCEPTION DES DEUX LISTES
noms_documents, ids_documents_rag = uploader_les_documents(DOSSIER_DOCUMENTS)
print(f"Noms uploadés : {noms_documents}")
print(f"IDs RAG pour suppression : {ids_documents_rag}\n")

# 2. UTILISATION DES NOMS DANS LE PROMPT INITIAL
# Convertit la liste de noms en une chaîne de caractères
noms_documents_str = ", ".join(noms_documents)

CHAPITRES_TITRES = [
    "Chapitre 1 : Identité du Projet",
    "Chapitre 2 : Architecture générale",
    "Chapitre 3 : Description étape par étape",
    "Chapitre 4 : Cartographie des Données (Technique)",
    "Chapitre 5 : Recommandations",
    "Chapitre 6 : Maintenabilité",
]

# PROMPT INITIAL DE CONTEXTE GLOBAL
# Le modèle est censé conserver ce contexte pour les requêtes suivantes.
PROMPT_INITIAL = f"""
Tu es un agent IA intégré dans Star Command. Ta mission est de générer une documentation technique vulgarisée
à partir des fichiers sources que je t'ai fournis, dont les IDs RAG sont : {noms_documents}.
Suis impérativement les règles et le format demandés. Nous allons procéder chapitre par chapitre.
"""

# PROMPTS SIMPLIFIÉS ET PLUS EFFICACES
# Le contexte (rôle, mission, IDs RAG) est géré par PROMPT_INITIAL, les prompts de chapitre sont directs.
# NOTE: J'ai laissé les titres explicites dans le f-string pour plus de clarté
PROMPTS = [
    f"Génère le chapitre '{titre}' en utilisant les documents {noms_documents}"
    for titre in CHAPITRES_TITRES
]


doc_word = Document()
doc_word = Document()

# 1. Ajout du titre principal
doc_word.add_heading("Documentation Technique - StarDoc✨", level=0)

# 2. Ajout de la phrase de mention légale en petit
phrase_ia = doc_word.add_paragraph()
run = phrase_ia.add_run(
    "Brouillon de documentation générée par l'IA sur la base des documents fournis dans la fiche StarCommand. Pour remplacer les diagrammes coller le code Mermaid dans https://app.diagrams.net/ > Organiser > Insérer > Mermaid"
)
run.font.size = Pt(9)  # Définit la taille à 9 points (petit)
run.italic = True  # Optionnel : mettre en italique pour le style
phrase_ia.alignment = 0  # Alignement à gauche (ou 1 pour centré)

# Pause de 4 minutes pour la vectorisation
print("Pause de 4 minutes en cours pour vectorisation RAG...")
time.sleep(240)
# ----------------------------------------------------
# 1. ENVOI DU CONTEXTE INITIAL
# ----------------------------------------------------
print("--- Envoi du contexte initial à l'agent ---")
interroger_agent_sse(
    PROMPT_INITIAL
)  # On envoie le contexte, mais on n'ajoute pas la réponse au document Word.
time.sleep(
    5
)  # Petite pause pour laisser le temps au contexte d'être potentiellement pris en compte par l'agent.

# ----------------------------------------------------
# 2.  GÉNÉRATION DES CHAPITRES
# ----------------------------------------------------


print("--- Étape 2 : Génération des chapitres ---")

for i, (prompt, titre_chapitre) in enumerate(zip(PROMPTS, CHAPITRES_TITRES), start=1):
    reponse = interroger_agent_sse(prompt)

    # --- Restitution du chapitre dans le terminal ---
    print("\n" + "=" * 80)
    print("=" * 80 + "\n")
    print(reponse)
    print("\n" + "=" * 80 + "\n")

    doc_word.add_heading(titre_chapitre, level=1)
    ajouter_texte_markdown(doc_word, reponse)
    doc_word.add_page_break()
    347
    print(f"📘 {titre_chapitre} généré.\n")

# ----------------------------------------------------
# 3. SAUVEGARDE ET NETTOYAGE
# ----------------------------------------------------
chemin_sortie = os.path.join(DOSSIER_DOCUMENTS, NOM_FICHIER_SORTIE)
doc_word.save(chemin_sortie)
print(f"✅ Document Word généré : {chemin_sortie}\n")

supprimer_documents(ids_documents_rag)

# --- Suppression des JSON générés ---
for f in os.listdir(DOSSIER_DOCUMENTS):
    if f.endswith("_parsed.json"):
        try:
            os.remove(os.path.join(DOSSIER_DOCUMENTS, f))
            print(f"🗑 JSON supprimé : {f}")
        except Exception as e:
            print(f"⚠ Impossible de supprimer {f} : {e}")

print("🎉 Script terminé.")
