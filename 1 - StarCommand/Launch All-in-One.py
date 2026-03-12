import configparser
import io
import json
import os
import re
import sys
import time
import tkinter as tk
import uuid
import xml.etree.ElementTree as ET
import zipfile
from tkinter import messagebox, simpledialog

import oletools
import requests
import sseclient
from docx import Document
from docx.shared import Pt
from openpyxl import load_workbook

# ==========================================================
# CHARGEMENT DE LA CONFIGURATION (config.ini)
# ==========================================================
config = configparser.ConfigParser()

if getattr(sys, "frozen", False):
    app_path = os.path.dirname(sys.executable)
else:
    app_path = os.path.dirname(__file__)

config_path = os.path.join(app_path, "config.ini")

if not os.path.exists(config_path):
    config["EDF_API"] = {"api_key": "CLE_ICI", "project_id": "ID_ICI"}
    config["PATHS"] = {"base_path": "CHEMIN_ICI"}
    config["SETTINGS"] = {"max_retry": "4", "wait_between_retry": "90"}
    with open(config_path, "w") as f:
        config.write(f)
    print(f"Fichier config.ini créé. Veuillez le remplir et relancer.")
    sys.exit()

config.read(config_path, encoding="utf-8")

API_KEY = config["EDF_API"]["api_key"]
PROJECT_ID = config["EDF_API"]["project_id"]
relative_base = config["PATHS"]["base_path"]
BASE_PATH = os.path.abspath(os.path.join(app_path, relative_base))

MAX_RETRY = int(config["SETTINGS"]["max_retry"])
WAIT_BETWEEN_RETRY = int(config["SETTINGS"]["wait_between_retry"])

print(f"📍 Chemin racine utilisé : {BASE_PATH}")


# ==========================================================
# INTERFACE DE SAISIE "UN CLIC"
# ==========================================================
def obtenir_num_dossier():
    root = tk.Tk()
    root.withdraw()
    num = simpledialog.askstring(
        "StarDoc ✨", "Veuillez saisir le numéro du dossier (ex: 378) :"
    )
    if not num:
        sys.exit()
    return num


NUM_DOSSIER = obtenir_num_dossier()
DOSSIER_DOCUMENTS = os.path.join(BASE_PATH, NUM_DOSSIER)

if not os.path.exists(DOSSIER_DOCUMENTS):
    messagebox.showerror("Erreur", f"Dossier introuvable :\n{DOSSIER_DOCUMENTS}")
    sys.exit()

# Session ID unique pour toute l'exécution
SESSION_ID = f"script_python_doc_gen_{uuid.uuid4().hex}"
print(f"🔑 Session ID : {SESSION_ID}")

# Paramètres API
URL_QUERY = "https://api.iag.edf.fr/v2/workspaces/HcA-puQ/webhooks/query"
HEADERS = {"knowledge-project-apikey": API_KEY}
NOM_FICHIER_SORTIE = f"STARDOC-{NUM_DOSSIER}.docx"


# ==========================================================
# UTILITIES (retry)
# ==========================================================
def requete_avec_retry(method, url, **kwargs):
    for tentative in range(1, MAX_RETRY + 1):
        try:
            response = requests.request(method, url, timeout=180, **kwargs)
            if response.status_code in [200, 201]:
                return response
            else:
                print(
                    f"⚠ Tentative {tentative}/{MAX_RETRY} - Code {response.status_code}"
                )
                if response.status_code == 504:
                    print(f"⏳ Attente {WAIT_BETWEEN_RETRY}s...")
                    time.sleep(WAIT_BETWEEN_RETRY)
        except requests.exceptions.Timeout:
            print(f"⏳ Timeout - tentative {tentative}")
            time.sleep(WAIT_BETWEEN_RETRY)
        except Exception as e:
            print(f"❌ Erreur : {e}")
            time.sleep(WAIT_BETWEEN_RETRY)
    return None


# ==========================================================
# FONCTIONS DE PARSING DE FICHIERS
# ==========================================================
def parse_uipath_xaml(file_path):
    try:
        tree = ET.parse(file_path)
        root = tree.getroot()
        activities, variables = [], []
        for elem in root.iter():
            tag = elem.tag.split("}")[-1]
            if "DisplayName" in elem.attrib:
                activities.append({"activity": tag, "name": elem.attrib["DisplayName"]})
            if "x:Name" in elem.attrib or "Name" in elem.attrib:
                variables.append(
                    {
                        "name": elem.attrib.get("x:Name", elem.attrib.get("Name")),
                        "type": tag,
                    }
                )
        return {"type": "UiPath", "activities": activities, "variables": variables}
    except Exception as e:
        return {"error": str(e)}


def parse_power_automate_json(file_path):
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            data = json.load(f)
        actions = data.get("definition", {}).get("actions", {})
        steps = [
            {"name": k, "type": v.get("type"), "inputs": v.get("inputs", {})}
            for k, v in actions.items()
        ]
        return {"type": "Power Automate", "steps": steps}
    except Exception as e:
        return {"error": str(e)}


def parse_powerbi_pbix(file_path):
    try:
        with zipfile.ZipFile(file_path, "r") as z:
            info = {
                "type": "Power BI (Deep Parse)",
                "tables_found": [],
                "measures_hint": [],
            }
            if "DataModel" in z.namelist():
                content = z.read("DataModel")
                strings = re.findall(rb"[A-Z][a-zA-Z0-9_]{2,30}", content)
                unique_strings = sorted(
                    list(set([s.decode("utf-8", errors="ignore") for s in strings]))
                )
                info["data_model_vocabulary"] = unique_strings[:200]

            if "Report/Layout" in z.namelist():
                layout = z.read("Report/Layout").decode("utf-16", errors="ignore")
                visuals = re.findall(r'"name":"([^"]+)"', layout)
                info["visual_elements"] = list(set(visuals))[:50]

            return info
    except Exception as e:
        return {"error": f"PBIX Parse Error: {str(e)}"}


def extract_vba_from_xlsm(file_path):
    try:
        with zipfile.ZipFile(file_path, "r") as z:
            namelist = z.namelist()
            if "xl/vbaProject.bin" not in namelist:
                return {"has_vba": False, "modules": []}

            vba_data = z.read("xl/vbaProject.bin")

            try:
                from oletools.olevba import VBA_Parser

                vba = VBA_Parser(filename="", data=vba_data)
                modules = []
                for (
                    filename,
                    stream_path,
                    vba_filename,
                    vba_code,
                ) in vba.extract_macros():
                    modules.append(
                        {
                            "module": vba_filename or filename,
                            "stream_path": " / ".join(stream_path)
                            if stream_path
                            else None,
                            "code": vba_code,
                        }
                    )
                vba.close()
                return {"has_vba": True, "modules": modules}
            except ImportError:
                return {
                    "has_vba": True,
                    "modules": [],
                    "warning": "oletools non installé, impossible d'extraire le code VBA",
                }
            except Exception as e:
                return {
                    "has_vba": True,
                    "modules": [],
                    "error": f"Erreur lors de l'extraction VBA : {str(e)}",
                }
    except Exception as e:
        return {
            "has_vba": False,
            "modules": [],
            "error": f"Erreur d'accès au vbaProject.bin : {str(e)}",
        }


def extract_excel_logic_universal(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    logic_report = {
        "filename": os.path.basename(file_path),
        "type": f"Excel ({ext})",
        "macros": extract_vba_from_xlsm(file_path) if ext == ".xlsm" else None,
        "potential_source_paths": [],
        "power_query_m": [],
        "ms_queries": [],
        "sheet_structures": {},
    }

    try:
        with zipfile.ZipFile(file_path, "r") as z:
            for item in [f for f in z.namelist() if "customXml/itemData" in f]:
                data = z.read(item)
                if b"PK\x03\x04" in data:
                    start = data.find(b"PK\x03\x04")
                    try:
                        with zipfile.ZipFile(io.BytesIO(data[start:])) as sub_z:
                            if "Formulas/Section1.m" in sub_z.namelist():
                                m_code = sub_z.read("Formulas/Section1.m").decode(
                                    "utf-16-le", errors="ignore"
                                )
                                logic_report["power_query_m"].extend(
                                    re.findall(
                                        r"let.*?in.*?(?=\r|\n|$)",
                                        m_code,
                                        re.DOTALL | re.IGNORECASE,
                                    )
                                )
                    except:
                        pass

            if "xl/connections.xml" in z.namelist():
                with z.open("xl/connections.xml") as f:
                    tree = ET.parse(f)
                    for conn in tree.findall(".//{*}connection"):
                        db_pr = conn.find(".//{*}dbPr")
                        if db_pr is not None:
                            logic_report["ms_queries"].append(
                                {"name": conn.get("name"), "sql": db_pr.get("command")}
                            )

        wb = load_workbook(file_path, data_only=False, read_only=True)
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]

            header_row_idx = 0
            headers = []
            for row in ws.iter_rows(min_row=1, max_row=10):
                row_values = [str(cell.value) for cell in row if cell.value is not None]
                if len(row_values) > 1:
                    headers = row_values
                    header_row_idx = row[0].row
                    break

            if not headers:
                continue

            formulas = {}
            data_row_idx = header_row_idx + 1
            rows_generator = list(
                ws.iter_rows(min_row=data_row_idx, max_row=data_row_idx)
            )

            if rows_generator:
                data_row = rows_generator[0]
                for cell in data_row:
                    val = cell.value
                    if val and isinstance(val, str):
                        col_idx = cell.column - 1
                        header_name = (
                            headers[col_idx]
                            if col_idx < len(headers)
                            else f"Col_{cell.column}"
                        )
                        if val.startswith("="):
                            formulas[header_name] = val
                        elif "\\\\" in val or (":" in val and "\\" in val):
                            if val not in logic_report["potential_source_paths"]:
                                logic_report["potential_source_paths"].append(val)

            if headers or formulas:
                logic_report["sheet_structures"][sheet_name] = {
                    "header_detected_at_line": header_row_idx,
                    "columns": headers,
                    "formulas_samples": formulas,
                }
    except Exception as e:
        logic_report["error"] = f"Erreur lors de l'analyse : {str(e)}"

    return logic_report


def extract_powerquery_from_excel(file_path):
    try:
        queries = []
        with zipfile.ZipFile(file_path, "r") as z:
            for name in z.namelist():
                if name.startswith("customXml/") and name.endswith(".xml"):
                    with z.open(name) as f:
                        try:
                            tree = ET.parse(f)
                            root = tree.getroot()
                            for elem in root.iter():
                                if elem.text and (
                                    "let" in elem.text
                                    or "in" in elem.text
                                    or "#" in elem.text
                                ):
                                    queries.append(elem.text.strip())
                        except ET.ParseError:
                            continue
        return queries
    except Exception as e:
        return [f"Error extracting Power Query: {str(e)}"]


def parse_powerquery_m(file_path):
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            lines = f.readlines()
            steps = [line.strip() for line in lines if line.strip()]
        return {"type": "Power Query", "steps": steps}
    except Exception as e:
        return {"error": str(e)}


# ==========================================================
# CONVERSION + LECTURE DU CONTENU (SANS UPLOAD RAG)
# ==========================================================

EXTENSIONS_A_CONVERTIR = {
    ".xaml": parse_uipath_xaml,
    ".pbix": parse_powerbi_pbix,
    ".m": parse_powerquery_m,
    ".xlsm": extract_excel_logic_universal,
    ".xlsx": extract_excel_logic_universal,
}

EXTENSIONS_TEXTE_BRUT = {
    ".txt",
    ".md",
    ".csv",
    ".json",
    ".sql",
    ".py",
    ".js",
    ".xml",
    ".html",
    ".log",
}


def lire_contenu_fichier(chemin_fichier):
    ext = os.path.splitext(chemin_fichier)[1].lower()
    nom = os.path.basename(chemin_fichier)

    if ext in EXTENSIONS_A_CONVERTIR:
        print(f"🔄 Conversion : {nom}")
        parser = EXTENSIONS_A_CONVERTIR[ext]
        result = parser(chemin_fichier)
        return json.dumps(result, indent=2, ensure_ascii=False)

    if ext in EXTENSIONS_TEXTE_BRUT:
        print(f"📄 Lecture directe : {nom}")
        try:
            with open(chemin_fichier, "r", encoding="utf-8", errors="replace") as f:
                return f.read()
        except Exception as e:
            print(f"⚠ Impossible de lire {nom} : {e}")
            return None

    print(f"⏭ Ignoré (format non supporté) : {nom}")
    return None


def charger_contenu_dossier(dossier_path):
    print("--- Chargement des fichiers (sans upload RAG) ---")
    contenu_par_fichier = {}

    for nom_fichier in sorted(os.listdir(dossier_path)):
        chemin_complet = os.path.join(dossier_path, nom_fichier)
        if not os.path.isfile(chemin_complet):
            continue
        if nom_fichier.startswith("STARDOC-"):
            continue

        contenu = lire_contenu_fichier(chemin_complet)
        if contenu:
            contenu_par_fichier[nom_fichier] = contenu
            print(f"✅ Chargé : {nom_fichier} ({len(contenu)} caractères)")

    print(f"--- {len(contenu_par_fichier)} fichier(s) chargé(s) ---\n")
    return contenu_par_fichier


def construire_bloc_contexte(contenu_par_fichier):
    blocs = []
    for nom, contenu in contenu_par_fichier.items():
        blocs.append(f"=== FICHIER : {nom} ===\n{contenu}\n=== FIN DE {nom} ===")
    return "\n\n".join(blocs)


# ==========================================================
# MOTEUR MARKDOWN → WORD
# ==========================================================
def clean_html_noise(text):
    if not text:
        return ""
    text = re.sub(r"<pr-canvas[^>]*>", "", text)
    text = re.sub(r"</pr-canvas>", "", text)
    text = re.sub(r"<[^>]+>", "", text)
    text = text.replace("**", "")
    return text


def sanitize_markdown(texte):
    """
    Nettoie :
    - Les séparateurs de tableaux corrompus (ex: | :----...---- |)
    - Les balises <br> HTML dans les cellules (remplacées par ' / ')
    - FIX: Clôture les lignes de tableau tronquées (cellule non fermée par |)
    """
    texte = re.sub(r"<br\s*/?>\s*", " / ", texte, flags=re.IGNORECASE)
    lignes_propres = []
    for ligne in texte.split("\n"):
        # FIX: clôturer les cellules de tableau non fermées (réponse tronquée)
        if "|" in ligne and ligne.count("|") >= 2:
            if not ligne.rstrip().endswith("|"):
                ligne = ligne.rstrip() + " |"
        # Normaliser les séparateurs de tableau trop longs
        if re.match(r"^\|?(\s*:?-{10,}:?\s*\|)+\s*$", ligne.strip()):
            nb_cols = max(ligne.count("|") - 1, 1)
            ligne = "| " + " | ".join(["---"] * nb_cols) + " |"
        lignes_propres.append(ligne)
    return "\n".join(lignes_propres)


def convert_markdown_table(table_text):
    rows = [r.strip() for r in table_text.strip().split("\n") if "|" in r]
    rows = [r.strip("|") for r in rows]
    rows = [r.split("|") for r in rows]

    if len(rows) < 2:
        return [], []

    header = rows[0]
    data = [r for r in rows[2:]]
    return header, data


def add_markdown_table_to_doc(doc, header, data):
    if not header and not data:
        return

    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    for i, text in enumerate(header):
        hdr_cells[i].text = text.replace("**", "").replace("`", "").strip()

    for row_data in data:
        row_cells = table.add_row().cells
        for i, text in enumerate(row_data):
            clean_text = text.replace("**", "").replace("`", "").strip()
            row_cells[i].text = clean_text

    doc.add_paragraph()


def ajouter_texte_markdown(doc, texte):
    texte = sanitize_markdown(texte)
    texte = texte.replace("\xa0", " ").replace("\u200b", "").strip()
    lignes = texte.split("\n")
    buffer_table = []

    def clean_stars(t):
        return t.replace("**", "").replace("*", "").strip()

    for ligne in lignes:
        ligne_raw = ligne.strip()
        if not ligne_raw:
            continue

        if "|" in ligne_raw and not ligne_raw.startswith("#"):
            buffer_table.append(ligne_raw)
            continue
        else:
            if buffer_table:
                header, data = convert_markdown_table("\n".join(buffer_table))
                add_markdown_table_to_doc(doc, header, data)
                buffer_table = []

        if ligne_raw.startswith("#"):
            titre = clean_stars(ligne_raw.lstrip("#"))
            doc.add_heading(titre, level=2)
            continue

        match_titre_section = re.match(r"^\s*[-*]\s+\*\*(.*?)\*\*\s*$", ligne_raw)
        if match_titre_section:
            clean_title = clean_stars(match_titre_section.group(1))
            doc.add_paragraph()
            p = doc.add_paragraph()
            run = p.add_run(clean_title)
            run.bold = True
            run.font.size = Pt(13)
            doc.add_paragraph()
            continue

        if ligne_raw.startswith("- ") or ligne_raw.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            texte_puce = clean_stars(ligne_raw.lstrip("-* "))
            p.add_run(texte_puce)
            continue

        p = doc.add_paragraph()
        p.add_run(clean_stars(ligne_raw))

    if buffer_table:
        header, data = convert_markdown_table("\n".join(buffer_table))
        add_markdown_table_to_doc(doc, header, data)


# ==========================================================
# INTERROGER L'AGENT EDF (SSE)
# FIX: Détection du finish_reason "length" → réponse tronquée rejetée
# ==========================================================
def interroger_agent_sse(question, max_tentatives=3):
    print(f"➡ Question (Mode SSE) : {question[:120]}...")

    body = {
        "text": question,
        "projectId": PROJECT_ID,
        "userId": SESSION_ID,
        "sse": True,
        "projectConfigOverride": {"ai": {"history": False}},
    }

    for tentative in range(1, max_tentatives + 1):
        reponse_complete = ""
        tous_les_events = []
        finish_reason = ""

        try:
            response = requests.post(
                URL_QUERY, headers=HEADERS, json=body, stream=True, timeout=300
            )

            if response.status_code != 200:
                print(f"⚠ Tentative {tentative} - Code {response.status_code}")
                log(
                    f"ERREUR HTTP - Tentative {tentative}",
                    f"Status : {response.status_code}\nBody envoyé :\n{json.dumps(body, indent=2, ensure_ascii=False)}",
                )
                time.sleep(WAIT_BETWEEN_RETRY)
                continue

            for event in sseclient.SSEClient(response).events():
                if not event.data:
                    continue
                tous_les_events.append(event.data)
                try:
                    data_json = json.loads(event.data)
                    if "answer" in data_json and data_json["answer"].strip():
                        reponse_complete = data_json["answer"]
                    # FIX: capturer le finish_reason pour détecter les troncatures
                    reasons = data_json.get("finishReasons", [])
                    if reasons:
                        finish_reason = reasons[0]
                except Exception:
                    continue

        except requests.exceptions.Timeout:
            print(f"⏳ Timeout - tentative {tentative}/{max_tentatives}")
            log(f"TIMEOUT - Tentative {tentative}", f"Question : {question[:300]}")
            time.sleep(WAIT_BETWEEN_RETRY)
        except Exception as e:
            print(f"❌ Erreur : {e} - tentative {tentative}/{max_tentatives}")
            log(
                f"EXCEPTION - Tentative {tentative}",
                f"Erreur : {e}\nQuestion : {question[:300]}",
            )
            time.sleep(10)

        # Log systématique de chaque tentative
        log(f"PROMPT - Tentative {tentative}/{max_tentatives}", question)
        log(
            f"FINISH REASON - Tentative {tentative}/{max_tentatives}",
            finish_reason if finish_reason else "(non détecté)",
        )
        log(
            f"EVENTS SSE BRUTS - Tentative {tentative}/{max_tentatives} ({len(tous_les_events)} events)",
            "\n---\n".join(tous_les_events)
            if tous_les_events
            else "(aucun event recu)",
        )
        log(
            f"REPONSE EXTRAITE - Tentative {tentative}/{max_tentatives} ({len(reponse_complete)} caracteres)",
            reponse_complete if reponse_complete else "(vide)",
        )

        # FIX: réponse tronquée par limite de tokens → invalide, on retry
        if finish_reason == "length":
            print(
                f"⚠ Réponse tronquée (finish_reason=length) — "
                f"tentative {tentative}/{max_tentatives}, retry avec consigne de concision..."
            )
            # Renforcer la consigne de concision dans le body pour le retry
            body["text"] = question + (
                "\n\nIMPORTANT : Ta réponse précédente était trop longue et a été coupée. "
                "Sois BEAUCOUP plus concis. MAX 60 caractères par cellule de tableau. "
                "Résume chaque étape en une seule phrase courte. "
                "Ne dépasse pas 15 lignes dans les tableaux."
            )
            time.sleep(10)
            continue

        if len(reponse_complete.strip()) >= 50:
            print(
                f"✅ Réponse reçue ({len(reponse_complete)} caractères, finish_reason={finish_reason})"
            )
            return reponse_complete

        print(
            f"⚠ Réponse insuffisante (tentative {tentative}/{max_tentatives}), retry..."
        )
        time.sleep(10)

    print("❌ Aucune réponse valide après toutes les tentatives.")
    return reponse_complete


# ==========================================================
# SCRIPT PRINCIPAL
# ==========================================================

# 1. CHARGEMENT DES FICHIERS
contenu_par_fichier = charger_contenu_dossier(DOSSIER_DOCUMENTS)

if not contenu_par_fichier:
    messagebox.showerror(
        "Erreur",
        "Aucun fichier lisible trouvé dans le dossier.\nVérifiez les formats supportés.",
    )
    sys.exit()

# 2. CONSTRUCTION DU BLOC DE CONTEXTE
bloc_contexte = construire_bloc_contexte(contenu_par_fichier)
noms_fichiers_str = ", ".join(contenu_par_fichier.keys())

print(f"📦 Taille totale du contexte injecté : {len(bloc_contexte)} caractères\n")

# 3. CHAPITRES À GÉNÉRER
CHAPITRES_TITRES = [
    "Chapitre 1 : Identité du Projet",
    "Chapitre 2 : Architecture générale",
    "Chapitre 3 : Description étape par étape (sous forme de tableau : Étape | Description | Outil utilisé)",
    "Chapitre 4 : Cartographie des Données (Technique)",
    "Chapitre 5 : Recommandations",
    "Chapitre 6 : Maintenabilité",
]

# 4. PROMPTS PAR CHAPITRE — contexte injecté dans chaque appel, historique désactivé
# FIX: consignes de concision renforcées sur les tableaux pour éviter finish_reason=length
PROMPTS = [
    f"""Tu es un agent IA générant une documentation technique vulgarisée.

Les fichiers analysés sont : {noms_fichiers_str}

Voici le contenu intégral de ces fichiers :

{bloc_contexte}

---
En te basant uniquement sur ces fichiers, génère le '{titre_chapitre}'.
Réponds directement avec le contenu du chapitre, sans introduction.

RÈGLES STRICTES DE FORMATAGE :
- MAX 120 caractères par cellule de tableau, sans exception.
- MAX 15 lignes dans les tableaux workflow.
- Si une description est longue, résume-la en une phrase courte.
- Utilise des bullet points séparés SOUS le tableau pour les détails si nécessaire.
- Ne mets jamais de longues phrases à l'intérieur d'une cellule de tableau."""
    for titre_chapitre in CHAPITRES_TITRES
]

# 5. INITIALISATION DU FICHIER DE LOG
LOG_PATH = os.path.join(DOSSIER_DOCUMENTS, f"STARDOC-DEBUG-{NUM_DOSSIER}.txt")


def log(titre, contenu):
    separateur = "=" * 80
    with open(LOG_PATH, "a", encoding="utf-8") as f:
        f.write(f"\n{separateur}\n{titre}\n{separateur}\n{contenu}\n")
    print(f"📝 Loggé : {titre}")


# Initialisation du fichier log
with open(LOG_PATH, "w", encoding="utf-8") as f:
    f.write(f"STARDOC DEBUG — Dossier {NUM_DOSSIER} — Session {SESSION_ID}\n")
    f.write(f"Date : {time.strftime('%Y-%m-%d %H:%M:%S')}\n")
    f.write(f"Fichiers chargés : {noms_fichiers_str}\n")
    f.write(f"Taille contexte : {len(bloc_contexte)} caractères\n")

# 6. INITIALISATION DU DOCUMENT WORD
doc_word = Document()
doc_word.add_heading("Documentation Technique - StarDoc✨", level=0)

phrase_ia = doc_word.add_paragraph()
run = phrase_ia.add_run(
    "Brouillon de documentation générée par l'IA sur la base des documents fournis dans la fiche StarCommand. "
    "Pour remplacer les diagrammes coller le code Mermaid dans https://app.diagrams.net/ > Organiser > Insérer > Mermaid"
)
run.font.size = Pt(9)
run.italic = True
phrase_ia.alignment = 0

# 7. GÉNÉRATION DES CHAPITRES
print("--- Génération des chapitres ---")

for i, (prompt, titre_chapitre) in enumerate(zip(PROMPTS, CHAPITRES_TITRES), start=1):
    reponse = interroger_agent_sse(prompt)

    print("\n" + "=" * 80)
    print("=" * 80 + "\n")
    print(reponse)
    print("\n" + "=" * 80 + "\n")

    doc_word.add_heading(titre_chapitre, level=1)
    ajouter_texte_markdown(doc_word, reponse)
    doc_word.add_page_break()

    print(f"📘 {titre_chapitre} généré.\n")

# 8. SAUVEGARDE
chemin_sortie = os.path.join(DOSSIER_DOCUMENTS, NOM_FICHIER_SORTIE)
doc_word.save(chemin_sortie)
print(f"✅ Document Word généré : {chemin_sortie}\n")

print("🎉 Script terminé.")
