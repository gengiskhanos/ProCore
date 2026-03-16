import configparser
import io
import json
import os
import re
import sys
import time
import tkinter as tk
import uuid
import xml.etree.ElementTree as ET
import zipfile
from datetime import datetime
from tkinter import messagebox, simpledialog

import oletools
import requests
import sseclient
from docx import Document
from docx.shared import Pt
from openpyxl import load_workbook

# ==========================================================
# VERSION & CHANGELOG
# ==========================================================
SCRIPT_VERSION = "v16-RAG-FIX"


# ==========================================================
# SYSTÈME DE LOGGING DEBUG
# ==========================================================
_logger_file = None
_log_path = None


def init_logger(dossier_projet):
    """Initialise le logger et crée le fichier de log dans le dossier projet."""
    global _logger_file, _log_path
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    _log_path = os.path.join(dossier_projet, f"DEBUG_STARDOC_{timestamp}.log")
    _logger_file = open(_log_path, "w", encoding="utf-8")
    log("=" * 80)
    log(f"  STARDOC DEBUG LOG — {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
    log(f"  Version script : {SCRIPT_VERSION}")
    log("=" * 80)
    print(f"📋 Log de debug : {_log_path}")


def log(message, level="INFO"):
    """Écrit dans le fichier de log ET dans la console."""
    ts = datetime.now().strftime("%H:%M:%S.%f")[:-3]
    line = f"[{ts}] [{level}] {message}"
    print(line)
    if _logger_file:
        _logger_file.write(line + "\n")
        _logger_file.flush()


def log_separateur(titre=""):
    log("")
    log("─" * 70)
    if titre:
        log(f"  {titre}")
        log("─" * 70)
    log("")


def log_json(label, data):
    """Loggue un objet JSON formaté."""
    log(f"{label} :")
    try:
        lines = json.dumps(data, indent=2, ensure_ascii=False).split("\n")
        for line in lines:
            log(f"  {line}")
    except Exception:
        log(f"  [Impossible de sérialiser en JSON : {type(data)}]")


def log_reponse_sse(
    reponse_complete,
    sources=None,
    search_results=None,
    nb_events=0,
    nb_chunks=0,
    duree_s=0,
    titre_chapitre="",
):
    """Log structuré d'une réponse SSE pour diagnostiquer la qualité RAG."""
    log_separateur(f"RÉPONSE SSE — {titre_chapitre}")
    log(f"⏱  Durée de la requête     : {duree_s:.2f}s")
    log(f"📦 Nb events SSE reçus     : {nb_events}")
    log(f"✂️  Nb chunks de réponse    : {nb_chunks}")
    log(f"📏 Longueur réponse finale : {len(reponse_complete)} caractères")

    log("")
    log("── DIAGNOSTICS QUALITÉ ──────────────────────────────────────────────")

    mots = reponse_complete.split()
    nb_mots = len(mots)
    log(f"  Nb mots dans la réponse     : {nb_mots}")

    # v7 : diagnostic ratio mots/chars
    if len(reponse_complete) > 0:
        ratio = nb_mots / len(reponse_complete)
        log(f"  Ratio mots/chars            : {ratio:.4f}")
        if ratio < 0.02 and len(reponse_complete) > 1000:
            log(
                f"  ⚠️  ALERTE : Ratio très faible (<0.02) — probable boucle de formatage",
                "WARN",
            )

    if len(reponse_complete) < 100:
        log(
            "  ⚠️  ALERTE : Réponse très courte (<100 chars) — possible échec RAG ou timeout",
            "WARN",
        )

    patterns_indesirables = [
        "Bien reçu",
        "Affirmatif",
        "J'attends",
        "Ceci conclut",
        "Avec plaisir",
        "Certainement",
        "En tant qu'IA",
    ]
    for p in patterns_indesirables:
        if p.lower() in reponse_complete.lower():
            log(
                f"  ⚠️  ALERTE : Formule indésirable détectée : '{p}' → le prompt système est peut-être ignoré",
                "WARN",
            )

    patterns_hallucination = [
        "je ne sais pas",
        "information non disponible",
        "je n'ai pas accès",
        "je n'ai pas trouvé",
        "aucune information",
        "données insuffisantes",
    ]
    for p in patterns_hallucination:
        if p.lower() in reponse_complete.lower():
            log(
                f"  ℹ️  Réponse contient : '{p}' → le RAG n'a peut-être pas trouvé de chunks pertinents",
                "INFO",
            )

    nb_placeholders = reponse_complete.lower().count("donnée non identifiée")
    if nb_placeholders > 0:
        log(
            f"  ⚠️  ALERTE : {nb_placeholders} placeholder(s) '[Donnée non identifiée]' détecté(s)",
            "WARN",
        )

    has_table = "|" in reponse_complete
    has_bullet = "- " in reponse_complete or "* " in reponse_complete
    has_heading = "#" in reponse_complete
    has_mermaid = "mermaid" in reponse_complete.lower() or "graph " in reponse_complete
    log(
        f"  Structure détectée → Tableau: {has_table} | Puces: {has_bullet} | Titres: {has_heading} | Mermaid: {has_mermaid}"
    )

    log("")
    log("── SOURCES RAG ──────────────────────────────────────────────────────")
    # v7 : filtrer les None pour ne pas alarmer à tort
    sources_reelles = [s for s in (sources or []) if s is not None]
    if sources_reelles:
        log(f"  Nb sources retournées : {len(sources_reelles)}")
        for i, s in enumerate(sources_reelles):
            log(f"  [{i + 1}] {s}")
    elif sources and len(sources) > 0:
        log(
            f"  ℹ️  {len(sources)} sources retournées mais sans URL (upload JSON = normal)",
            "INFO",
        )
    else:
        log(
            "  ⚠️  Aucune source RAG retournée",
            "WARN",
        )

    if search_results:
        log("")
        log("── CHUNKS RAG RÉCUPÉRÉS ─────────────────────────────────────────────")
        log(f"  Nb chunks : {len(search_results)}")
        for i, chunk in enumerate(search_results):
            log(f"  ── Chunk [{i + 1}] ──")
            log(f"     Nom doc : {chunk.get('name', 'N/A')}")
            log(f"     ID      : {chunk.get('_id', 'N/A')}")
            content_preview = (chunk.get("content") or "")[:300].replace("\n", " ")
            log(f"     Contenu (300 chars) : {content_preview}")
    else:
        log(
            "  ⚠️  search.results absent de la réponse finale",
            "WARN",
        )

    log("")
    log("── RÉPONSE BRUTE (extrait) ──────────────────────────────────────────")
    if reponse_complete:
        log(f"  [DÉBUT] {reponse_complete[:500].replace(chr(10), ' ↵ ')}")
        if len(reponse_complete) > 700:
            log(f"  [FIN]   {reponse_complete[-200:].replace(chr(10), ' ↵ ')}")
    else:
        log("  [VIDE]", "WARN")
    log("")


# ==========================================================
# CHARGEMENT DE LA CONFIGURATION (config.ini)
# ==========================================================
config = configparser.ConfigParser()

if getattr(sys, "frozen", False):
    app_path = os.path.dirname(sys.executable)
else:
    app_path = os.path.dirname(__file__)

config_path = os.path.join(app_path, "config.ini")

if not os.path.exists(config_path):
    config["EDF_API"] = {"api_key": "CLE_ICI", "project_id": "ID_ICI"}
    config["PATHS"] = {"base_path": "CHEMIN_ICI"}
    config["SETTINGS"] = {"max_retry": "4", "wait_between_retry": "90"}
    with open(config_path, "w") as f:
        config.write(f)
    print(f"Fichier config.ini créé. Veuillez le remplir et relancer.")
    sys.exit()

config.read(config_path, encoding="utf-8")

API_KEY = config["EDF_API"]["api_key"]
PROJECT_ID = config["EDF_API"]["project_id"]
relative_base = config["PATHS"]["base_path"]
BASE_PATH = os.path.abspath(os.path.join(app_path, relative_base))
MAX_RETRY = int(config["SETTINGS"]["max_retry"])
WAIT_BETWEEN_RETRY = int(config["SETTINGS"]["wait_between_retry"])
RUN_ID = f"stardoc_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{uuid.uuid4().hex[:6]}"
print(f"📍 Chemin racine utilisé : {BASE_PATH}")
print(f"🔖 Version du script : {SCRIPT_VERSION}")
print(f"🔑 Id du run : {RUN_ID}")


# ==========================================================
# INTERFACE DE SAISIE "UN CLIC"
# ==========================================================
def obtenir_num_dossier():
    root = tk.Tk()
    root.withdraw()
    num = simpledialog.askstring(
        "StarDoc ✨", "Veuillez saisir le numéro du dossier (ex: 378) :"
    )
    if not num:
        sys.exit()
    return num


NUM_DOSSIER = obtenir_num_dossier()
DOSSIER_DOCUMENTS = os.path.join(BASE_PATH, NUM_DOSSIER)

if not os.path.exists(DOSSIER_DOCUMENTS):
    messagebox.showerror("Erreur", f"Dossier introuvable :\n{DOSSIER_DOCUMENTS}")
    sys.exit()

init_logger(DOSSIER_DOCUMENTS)

log_separateur("CONFIGURATION")
log(f"Dossier projet     : {DOSSIER_DOCUMENTS}")
log(f"NUM_DOSSIER        : {NUM_DOSSIER}")
log(f"PROJECT_ID         : {PROJECT_ID}")
log(f"API_KEY (masquée)  : {API_KEY[:6]}...{API_KEY[-4:]}")
log(f"MAX_RETRY          : {MAX_RETRY}")
log(f"WAIT_BETWEEN_RETRY : {WAIT_BETWEEN_RETRY}s")

# Paramètres API
URL_DOCUMENT = "https://api.iag.edf.fr/v2/workspaces/HcA-puQ/webhooks/document"
URL_QUERY = "https://api.iag.edf.fr/v2/workspaces/HcA-puQ/webhooks/query"
URL_DELETE = "https://api.iag.edf.fr/v2/workspaces/HcA-puQ/webhooks/document"
HEADERS = {"knowledge-project-apikey": API_KEY}
NOM_FICHIER_SORTIE = f"STARDOC-{NUM_DOSSIER}.docx"

# Durée de la pause de vectorisation (secondes)
PAUSE_VECTORISATION = 120


# ==========================================================
# UTILITIES (retry avec logging)
# ==========================================================
def requete_avec_retry(method, url, **kwargs):
    for tentative in range(1, MAX_RETRY + 1):
        log(f"🌐 [{method}] {url[:120]}... — Tentative {tentative}/{MAX_RETRY}")
        try:
            t0 = time.time()
            response = requests.request(method, url, timeout=180, **kwargs)
            duree = time.time() - t0
            log(f"   → HTTP {response.status_code} en {duree:.2f}s")

            if response.status_code in [200, 201]:
                return response
            else:
                log(
                    f"   ⚠️  Code inattendu {response.status_code} — body: {response.text[:300]}",
                    "WARN",
                )
                if response.status_code in [504, 429]:
                    log(f"   ⏳ Attente {WAIT_BETWEEN_RETRY}s...", "WARN")
                    time.sleep(WAIT_BETWEEN_RETRY)
                else:
                    time.sleep(5)
        except requests.exceptions.Timeout:
            log(f"   ⏳ Timeout réseau à la tentative {tentative}", "WARN")
            time.sleep(WAIT_BETWEEN_RETRY)
        except Exception as e:
            log(f"   ❌ Exception : {e}", "ERROR")
            time.sleep(WAIT_BETWEEN_RETRY)
    log(f"   ❌ Toutes les tentatives ont échoué pour [{method}] {url[:80]}", "ERROR")
    return None


# ==========================================================
# FONCTIONS DE PARSING DE FICHIERS (conversion pré-upload)
# ==========================================================
def parse_uipath_xaml(file_path):
    try:
        tree = ET.parse(file_path)
        root = tree.getroot()
        activities, variables = [], []
        for elem in root.iter():
            tag = elem.tag.split("}")[-1]
            if "DisplayName" in elem.attrib:
                activities.append({"activity": tag, "name": elem.attrib["DisplayName"]})
            if "x:Name" in elem.attrib or "Name" in elem.attrib:
                variables.append(
                    {
                        "name": elem.attrib.get("x:Name", elem.attrib.get("Name")),
                        "type": tag,
                    }
                )
        log(f"   XAML parsé : {len(activities)} activités, {len(variables)} variables")
        return {"type": "UiPath", "activities": activities, "variables": variables}
    except Exception as e:
        log(f"   ❌ Erreur parse XAML : {e}", "ERROR")
        return {"error": str(e)}


def parse_powerapps_yaml(file_path):
    """
    v7.1 : Parse les fichiers Power Apps Canvas (.pa.yaml).
    Le YAML Power Apps utilise '=' comme préfixe de formule PowerFx,
    ce qui casse yaml.safe_load. On parse donc en mode texte/regex.
    Extrait : écrans, contrôles, formules, navigations, variables, data sources.
    """
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()

        basename = os.path.basename(file_path)
        result = {"type": "Power Apps Canvas", "filename": basename}

        # --- EditorState ---
        if "EditorState:" in content:
            result["type"] = "Power Apps EditorState"
            result["screens_order"] = re.findall(r"^\s+- (\w+)$", content, re.MULTILINE)
            log(f"   Power Apps EditorState : {result['screens_order']}")
            return result

        # --- App config ---
        if content.strip().startswith("App:") or "\nApp:" in content:
            result["type"] = "Power Apps App Config"
            onstart = re.search(r"OnStart:\s*\|?-?\s*\n((?:\s{4,}.*\n)*)", content)
            if onstart:
                result["on_start"] = onstart.group(1).strip()
            start = re.search(r"StartScreen:\s*=(\w+)", content)
            if start:
                result["start_screen"] = start.group(1)
            log(f"   Power Apps App Config : start_screen={result.get('start_screen')}")
            return result

        # --- Screen file ---
        screen = re.search(r"Screens:\n\s+(\w+):", content)
        if screen:
            result["screen_name"] = screen.group(1)

        # Controls avec leur type
        controls = re.findall(r"- (\w+):\s*\n\s+Control:\s*(\S+)", content)
        result["controls"] = [{"name": c[0], "type": c[1]} for c in controls]

        # Navigate() calls
        result["navigations"] = sorted(set(re.findall(r"Navigate\((\w+)", content)))

        # Variables Set() et Collect()
        result["variables_set"] = sorted(set(re.findall(r"Set\((\w+)", content)))
        result["collections"] = sorted(
            set(re.findall(r"(?:Collect|ClearCollect)\((\w+)", content))
        )

        # Formules significatives (OnSelect, OnChange, OnVisible, Items)
        formulas = []
        for key, formula in re.findall(
            r"(OnSelect|OnChange|OnVisible|Items|OnStart):\s*=(.*)", content
        ):
            f = formula.strip()
            if len(f) > 10:
                formulas.append({"property": key, "formula": f[:300]})
        # Formules multilignes
        for m in re.finditer(
            r"(OnSelect|OnChange|OnVisible|Items):\s*\|?-?\s*\n(\s+=.*?)(?=\n\s+\w+:|\n\s+-\s)",
            content,
            re.DOTALL,
        ):
            f = m.group(2).strip().lstrip("=")
            if len(f) > 10:
                formulas.append({"property": m.group(1), "formula": f[:300]})
        result["formulas"] = formulas[:40]

        # Data sources — noms entre apostrophes dans les formules
        candidates = re.findall(r"'([A-Z][A-Za-zÀ-ÿ0-9_ -]{3,50})'", content)
        EXCLUDE = {
            "Font",
            "Segoe UI",
            "Open Sans",
            "RGBA",
            "Color",
            "Self",
            "Parent",
            "ThisItem",
            "Align",
            "VerticalAlign",
            "DisplayMode",
            "FontWeight",
            "ScreenTransition",
            "Icon",
            "BorderStyle",
            "DataSourceInfo",
        }
        ds = sorted(set(c for c in candidates if c not in EXCLUDE))
        result["data_sources"] = ds[:20]

        log(
            f"   Power Apps Screen '{result.get('screen_name')}' : "
            f"{len(result['controls'])} contrôles, "
            f"{len(result['formulas'])} formules, "
            f"navigations={result['navigations']}"
        )
        return result
    except Exception as e:
        log(f"   ❌ Erreur parse Power Apps YAML : {e}", "ERROR")
        return {"error": str(e)}


def parse_power_automate_json(file_path):
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            data = json.load(f)
        actions = data.get("definition", {}).get("actions", {})
        steps = [
            {"name": k, "type": v.get("type"), "inputs": v.get("inputs", {})}
            for k, v in actions.items()
        ]
        log(f"   Power Automate parsé : {len(steps)} étapes")
        return {"type": "Power Automate", "steps": steps}
    except Exception as e:
        log(f"   ❌ Erreur parse Power Automate JSON : {e}", "ERROR")
        return {"error": str(e)}


def parse_powerbi_pbix(file_path):
    try:
        with zipfile.ZipFile(file_path, "r") as z:
            info = {
                "type": "Power BI (Deep Parse)",
                "tables_found": [],
                "measures_hint": [],
            }
            if "DataModel" in z.namelist():
                content = z.read("DataModel")
                strings = re.findall(rb"[A-Z][a-zA-Z0-9_]{2,30}", content)
                unique_strings = sorted(
                    list(set([s.decode("utf-8", errors="ignore") for s in strings]))
                )
                info["data_model_vocabulary"] = unique_strings[:200]
                log(f"   PBIX DataModel : {len(unique_strings)} termes extraits")
            if "Report/Layout" in z.namelist():
                layout = z.read("Report/Layout").decode("utf-16", errors="ignore")
                visuals = re.findall(r'"name":"([^"]+)"', layout)
                info["visual_elements"] = list(set(visuals))[:50]
                log(f"   PBIX Layout : {len(info['visual_elements'])} visuels détectés")
            return info
    except Exception as e:
        log(f"   ❌ Erreur parse PBIX : {e}", "ERROR")
        return {"error": f"PBIX Parse Error: {str(e)}"}


def extract_vba_from_xlsm(file_path):
    try:
        with zipfile.ZipFile(file_path, "r") as z:
            namelist = z.namelist()
            if "xl/vbaProject.bin" not in namelist:
                return {"has_vba": False, "modules": []}
            vba_data = z.read("xl/vbaProject.bin")
            try:
                from oletools.olevba import VBA_Parser

                vba = VBA_Parser(filename="", data=vba_data)
                modules = []
                for (
                    filename,
                    stream_path,
                    vba_filename,
                    vba_code,
                ) in vba.extract_macros():
                    modules.append(
                        {
                            "module": vba_filename or filename,
                            "stream_path": " / ".join(stream_path)
                            if stream_path
                            else None,
                            "code": vba_code,
                        }
                    )
                vba.close()
                log(f"   VBA extrait : {len(modules)} modules")
                return {"has_vba": True, "modules": modules}
            except ImportError:
                log("   ⚠️  oletools non installé, VBA non extrait", "WARN")
                return {
                    "has_vba": True,
                    "modules": [],
                    "warning": "oletools non installé",
                }
            except Exception as e:
                log(f"   ⚠️  Erreur extraction VBA : {e}", "WARN")
                return {"has_vba": True, "modules": [], "error": str(e)}
    except Exception as e:
        log(f"   ❌ Erreur accès vbaProject.bin : {e}", "ERROR")
        return {"has_vba": False, "modules": [], "error": str(e)}


def extract_excel_logic_universal(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    logic_report = {
        "filename": os.path.basename(file_path),
        "type": f"Excel ({ext})",
        "macros": extract_vba_from_xlsm(file_path) if ext == ".xlsm" else None,
        "potential_source_paths": [],
        "power_query_m": [],
        "ms_queries": [],
        "sheet_structures": {},
    }
    try:
        with zipfile.ZipFile(file_path, "r") as z:
            for item in [f for f in z.namelist() if "customXml/itemData" in f]:
                data = z.read(item)
                if b"PK\x03\x04" in data:
                    start = data.find(b"PK\x03\x04")
                    try:
                        with zipfile.ZipFile(io.BytesIO(data[start:])) as sub_z:
                            if "Formulas/Section1.m" in sub_z.namelist():
                                m_code = sub_z.read("Formulas/Section1.m").decode(
                                    "utf-16-le", errors="ignore"
                                )
                                logic_report["power_query_m"].extend(
                                    re.findall(
                                        r"let.*?in.*?(?=\r|\n|$)",
                                        m_code,
                                        re.DOTALL | re.IGNORECASE,
                                    )
                                )
                    except:
                        pass
            if "xl/connections.xml" in z.namelist():
                with z.open("xl/connections.xml") as f:
                    tree = ET.parse(f)
                    for conn in tree.findall(".//{*}connection"):
                        db_pr = conn.find(".//{*}dbPr")
                        if db_pr is not None:
                            logic_report["ms_queries"].append(
                                {"name": conn.get("name"), "sql": db_pr.get("command")}
                            )
        wb = load_workbook(file_path, data_only=False, read_only=True)
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            header_row_idx = 0
            headers = []
            for row in ws.iter_rows(min_row=1, max_row=10):
                row_values = [str(cell.value) for cell in row if cell.value is not None]
                if len(row_values) > 1:
                    headers = row_values
                    header_row_idx = row[0].row
                    break
            if not headers:
                continue
            formulas = {}
            data_row_idx = header_row_idx + 1
            rows_generator = list(
                ws.iter_rows(min_row=data_row_idx, max_row=data_row_idx)
            )
            if rows_generator:
                data_row = rows_generator[0]
                for cell in data_row:
                    val = cell.value
                    if val and isinstance(val, str):
                        col_idx = cell.column - 1
                        header_name = (
                            headers[col_idx]
                            if col_idx < len(headers)
                            else f"Col_{cell.column}"
                        )
                        if val.startswith("="):
                            formulas[header_name] = val
                        elif "\\\\" in val or (":" in val and "\\" in val):
                            if val not in logic_report["potential_source_paths"]:
                                logic_report["potential_source_paths"].append(val)
            if headers or formulas:
                logic_report["sheet_structures"][sheet_name] = {
                    "header_detected_at_line": header_row_idx,
                    "columns": headers,
                    "formulas_samples": formulas,
                }
        log(
            f"   Excel analysé : {len(logic_report['sheet_structures'])} feuilles, "
            f"{len(logic_report['power_query_m'])} requêtes M, "
            f"{len(logic_report['ms_queries'])} connexions SQL"
        )
    except Exception as e:
        log(f"   ❌ Erreur analyse Excel : {e}", "ERROR")
        logic_report["error"] = str(e)
    return logic_report


def extract_powerquery_from_excel(file_path):
    try:
        queries = []
        with zipfile.ZipFile(file_path, "r") as z:
            for name in z.namelist():
                if name.startswith("customXml/") and name.endswith(".xml"):
                    with z.open(name) as f:
                        try:
                            tree = ET.parse(f)
                            root = tree.getroot()
                            for elem in root.iter():
                                if elem.text and (
                                    "let" in elem.text
                                    or "in" in elem.text
                                    or "#" in elem.text
                                ):
                                    queries.append(elem.text.strip())
                        except ET.ParseError:
                            continue
        return queries
    except Exception as e:
        return [f"Error extracting Power Query: {str(e)}"]


def parse_powerquery_m(file_path):
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            lines = f.readlines()
            steps = [line.strip() for line in lines if line.strip()]
        return {"type": "Power Query", "steps": steps}
    except Exception as e:
        return {"error": str(e)}


# ==========================================================
# CONVERSION AUTOMATIQUE AVANT UPLOAD
# ==========================================================
def convertir_si_necessaire(chemin_fichier):
    ext = os.path.splitext(chemin_fichier)[1].lower()
    mapping = {
        ".xaml": parse_uipath_xaml,
        ".pbix": parse_powerbi_pbix,
        ".m": parse_powerquery_m,
        ".xlsm": extract_excel_logic_universal,
        ".xlsx": extract_excel_logic_universal,
        ".yaml": parse_powerapps_yaml,
        ".yml": parse_powerapps_yaml,
    }
    if ext not in mapping:
        return chemin_fichier
    log(f"🔄 Conversion de {os.path.basename(chemin_fichier)} (ext: {ext})")
    parser = mapping[ext]
    result = parser(chemin_fichier)
    chemin_json = chemin_fichier + "_parsed.json"
    with open(chemin_json, "w", encoding="utf-8") as f:
        json.dump(result, f, indent=2, ensure_ascii=False)
    taille = os.path.getsize(chemin_json)
    log(f"   → JSON généré : {os.path.basename(chemin_json)} ({taille} octets)")
    return chemin_json


# ==========================================================
# EXTRACTION DOCX ENRICHIE (paragraphes + tableaux)
# ==========================================================
def extraire_texte_docx_complet(chemin_docx):
    """
    Extrait le texte complet d'un fichier .docx en incluant
    les paragraphes ET les tableaux (que python-docx ne met
    pas dans doc.paragraphs).
    """
    try:
        from docx import Document as DocxReader

        doc_temp = DocxReader(chemin_docx)
        texte_parties = []

        for p in doc_temp.paragraphs:
            if p.text.strip():
                texte_parties.append(p.text.strip())

        for table in doc_temp.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.replace("|", "").replace(" ", ""):
                    texte_parties.append(row_text)

        texte_complet = "\n".join(texte_parties)
        log(
            f"   📄 Texte extrait du .docx : {len(texte_complet)} chars "
            f"({len([p for p in doc_temp.paragraphs if p.text.strip()])} paragraphes, "
            f"{len(doc_temp.tables)} tableaux)"
        )
        return texte_complet
    except Exception as e:
        log(f"   ⚠️  Impossible d'extraire le texte du .docx : {e}", "WARN")
        return None


# ==========================================================
# MOTEUR MARKDOWN → WORD
# ==========================================================
def clean_html_noise(text):
    if not text:
        return ""
    text = re.sub(r"<pr-canvas[^>]*>", "", text)
    text = re.sub(r"</pr-canvas>", "", text)
    text = re.sub(r"<[^>]+>", "", text)
    text = text.replace("**", "")
    return text


def convert_markdown_table(table_text):
    rows = [r.strip() for r in table_text.strip().split("\n") if "|" in r]
    rows = [r.strip("|") for r in rows]
    rows = [r.split("|") for r in rows]
    if len(rows) < 2:
        return [], []
    header = rows[0]
    data = [r for r in rows[2:]]
    return header, data


def add_markdown_table_to_doc(doc, header, data):
    if not header and not data:
        return
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, text in enumerate(header):
        hdr_cells[i].text = text.replace("**", "").replace("`", "").strip()
    for row_data in data:
        row_cells = table.add_row().cells
        for i, text in enumerate(row_data):
            clean_text = text.replace("**", "").replace("`", "").strip()
            row_cells[i].text = clean_text
    doc.add_paragraph()


def ajouter_texte_markdown(doc, texte):
    texte = texte.replace("\xa0", " ").replace("\u200b", "").strip()
    lignes = texte.split("\n")
    buffer_table = []

    def clean_stars(t):
        return t.replace("**", "").replace("*", "").strip()

    for ligne in lignes:
        ligne_raw = ligne.strip()
        if not ligne_raw:
            continue
        if "|" in ligne_raw and not ligne_raw.startswith("#"):
            buffer_table.append(ligne_raw)
            continue
        else:
            if buffer_table:
                header, data = convert_markdown_table("\n".join(buffer_table))
                add_markdown_table_to_doc(doc, header, data)
                buffer_table = []
        if ligne_raw.startswith("#"):
            titre = clean_stars(ligne_raw.lstrip("#"))
            doc.add_heading(titre, level=2)
            continue
        match_titre_section = re.match(r"^\s*[-*]\s+\*\*(.*?)\*\*\s*$", ligne_raw)
        if match_titre_section:
            clean_title = clean_stars(match_titre_section.group(1))
            doc.add_paragraph()
            p = doc.add_paragraph()
            run = p.add_run(clean_title)
            run.bold = True
            run.font.size = Pt(13)
            doc.add_paragraph()
            continue
        if ligne_raw.startswith("- ") or ligne_raw.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            texte_puce = clean_stars(ligne_raw.lstrip("-* "))
            p.add_run(texte_puce)
            continue
        p = doc.add_paragraph()
        p.add_run(clean_stars(ligne_raw))

    if buffer_table:
        header, data = convert_markdown_table("\n".join(buffer_table))
        add_markdown_table_to_doc(doc, header, data)


# ==========================================================
# UPLOAD DES DOCUMENTS (avec extraction docx enrichie)
# ==========================================================
def uploader_les_documents(dossier_path):
    log_separateur("ÉTAPE 1 : UPLOAD DES FICHIERS VERS LE RAG")
    noms_documents_uploades = []
    ids_documents_rag = []

    EXTENSIONS_EXCLUES = {".log"}
    PREFIXES_EXCLUS = ("DEBUG_STARDOC", "STARDOC-")

    tous_fichiers = [
        f
        for f in os.listdir(dossier_path)
        if os.path.isfile(os.path.join(dossier_path, f))
    ]
    fichiers = []
    for f in tous_fichiers:
        ext = os.path.splitext(f)[1].lower()
        if ext in EXTENSIONS_EXCLUES:
            log(f"  ⏭  Ignoré (extension exclue) : {f}")
            continue
        if any(f.startswith(p) for p in PREFIXES_EXCLUS):
            log(f"  ⏭  Ignoré (fichier généré StarDoc) : {f}")
            continue
        fichiers.append(f)

    log(
        f"Fichiers détectés dans le dossier : {len(tous_fichiers)} total, {len(fichiers)} à uploader"
    )
    for f in fichiers:
        taille = os.path.getsize(os.path.join(dossier_path, f))
        log(f"  - {f} ({taille} octets)")
    log("")

    for nom_fichier in fichiers:
        chemin_complet = os.path.join(dossier_path, nom_fichier)
        log(f"── Traitement : {nom_fichier}")

        chemin_upload = convertir_si_necessaire(chemin_complet)
        nom_upload = os.path.basename(chemin_upload)
        taille_upload = os.path.getsize(chemin_upload)
        log(f"   Fichier à uploader : {nom_upload} ({taille_upload} octets)")

        t0 = time.time()

        ext_upload = os.path.splitext(chemin_upload)[1].lower()

        texte_contenu = None
        if ext_upload in [".json", ".txt", ".csv", ".m", ".py", ".xml"]:
            try:
                with open(chemin_upload, "r", encoding="utf-8") as f:
                    texte_contenu = f.read()
            except UnicodeDecodeError:
                with open(chemin_upload, "r", encoding="latin-1") as f:
                    texte_contenu = f.read()
        elif ext_upload in [".docx"]:
            texte_contenu = extraire_texte_docx_complet(chemin_upload)

        if texte_contenu:
            json_body = {
                "projectId": PROJECT_ID,
                "name": nom_upload,
                "content": {"text": texte_contenu},
                "tags": [],
            }
            headers_json = {**HEADERS, "Content-Type": "application/json"}
            log(f"   📤 Mode upload : JSON (content.text) — {len(texte_contenu)} chars")
            response = requete_avec_retry(
                "POST", URL_DOCUMENT, headers=headers_json, json=json_body
            )
        else:
            log(f"   ℹ️  Fichier binaire — upload en multipart (sans tags)")
            with open(chemin_upload, "rb") as f:
                files = {"file": (nom_upload, f)}
                data = {"projectId": PROJECT_ID, "name": nom_upload}
                response = requete_avec_retry(
                    "POST", URL_DOCUMENT, headers=HEADERS, data=data, files=files
                )

        duree_upload = time.time() - t0

        doc_id = None
        if response:
            try:
                resp_json = response.json()
                log_json(f"   Réponse API upload de {nom_upload}", resp_json)

                # v7 : chercher l'ID dans plusieurs clés possibles
                doc_id = resp_json.get("_id") or resp_json.get("id")
                if (
                    not doc_id
                    and "search" in resp_json
                    and "results" in resp_json["search"]
                ):
                    results = resp_json["search"]["results"]
                    if results:
                        doc_id = results[0].get("_id") or results[0].get("id")
            except Exception as e:
                log(f"   ❌ Erreur parsing réponse upload : {e}", "ERROR")

            if doc_id:
                noms_documents_uploades.append(nom_upload)
                ids_documents_rag.append(doc_id)
                log(
                    f"   ✅ Upload OK — ID RAG : {doc_id} — durée : {duree_upload:.2f}s"
                )
            else:
                log(f"   ⚠️  ID manquant dans la réponse", "WARN")
        else:
            log(f"   ❌ Échec total de l'upload pour {nom_upload}", "ERROR")
        log("")

    log(
        f"Upload terminé — {len(noms_documents_uploades)}/{len(fichiers)} fichiers indexés"
    )
    log(f"IDs RAG : {ids_documents_rag}")
    return noms_documents_uploades, ids_documents_rag


# ==========================================================


# ==========================================================
# NORMALISATION MARKDOWN — v7 anti-boucle
# ==========================================================
def normaliser_tableaux_markdown(texte):
    """
    v7 : Normalise les séparateurs de tableau Markdown.
    Remplace les longues séries de tirets (cause de boucle Gemini)
    par des séparateurs courts et propres.
    """
    if not texte:
        return texte

    lignes = texte.split("\n")
    resultat = []

    for ligne in lignes:
        stripped = ligne.strip()
        # Détecter les lignes de séparateurs de tableau : |---|---|...|
        if re.match(r"^\|[\s\-:|]+\|$", stripped):
            # Compter le nombre de colonnes
            nb_cols = stripped.count("|") - 1
            if nb_cols > 0:
                separateur_propre = "| " + " | ".join(["---"] * nb_cols) + " |"
                resultat.append(separateur_propre)
            else:
                resultat.append(stripped)
        else:
            resultat.append(ligne)

    return "\n".join(resultat)


# ==========================================================
# DÉTECTION DE BOUCLE DE RÉPÉTITION — v7 améliorée
# ==========================================================
def nettoyer_reponse_boucle(reponse):
    if not reponse or len(reponse) < 3000:
        return reponse

    if "InternalServerError" in reponse:
        idx = reponse.find("litellm.InternalServerError")
        if idx == -1:
            idx = reponse.find("InternalServerError")
        if idx > 0:
            reponse = reponse[:idx].rstrip()
            log("   🔧 Message d'erreur InternalServerError retiré", "WARN")

    # Détection de lignes de tirets répétées (boucle de séparateurs tableau)
    lignes = reponse.split("\n")
    lignes_nettoyees = []
    compteur_tirets = 0
    for ligne in lignes:
        contenu_sans_separateurs = (
            ligne.replace("-", "").replace("|", "").replace(":", "").replace(" ", "")
        )
        if len(ligne) > 20 and len(contenu_sans_separateurs) < 5:
            compteur_tirets += 1
            if compteur_tirets <= 3:
                lignes_nettoyees.append(ligne)
        else:
            compteur_tirets = 0
            lignes_nettoyees.append(ligne)

    reponse_nettoyee = "\n".join(lignes_nettoyees)

    # Détection de blocs de texte répétés
    if len(reponse_nettoyee) > 8000:
        moitie = len(reponse_nettoyee) // 2
        premiere_moitie = reponse_nettoyee[:moitie]
        deuxieme_moitie = reponse_nettoyee[moitie:]
        sample_1 = premiere_moitie[:500].strip()
        sample_2 = deuxieme_moitie[:500].strip()
        if sample_1 and sample_2:
            if sample_2[:200] in premiere_moitie:
                reponse_nettoyee = premiere_moitie.rstrip()
                log(
                    f"   🔧 Répétition de bloc détectée — tronqué à la 1re moitié ({len(reponse_nettoyee)} chars)",
                    "WARN",
                )

    if len(reponse_nettoyee) < len(reponse) * 0.8:
        log(
            f"   🔧 Boucle nettoyée : {len(reponse)} → {len(reponse_nettoyee)} chars",
            "WARN",
        )

    # v7 : normaliser les séparateurs de tableau après nettoyage
    reponse_nettoyee = normaliser_tableaux_markdown(reponse_nettoyee)

    return reponse_nettoyee


# ==========================================================
# VALIDATION QUALITÉ — v7 nouveau
# ==========================================================
def valider_qualite_chapitre(titre, reponse):
    """
    v7.1 : Valide la qualité d'une réponse générée.
    Retourne (True, "OK") si exploitable, (False, raison) sinon.
    Inclut des seuils minimum par chapitre pour les chapitres multi-sections.
    """
    issues = []

    if not reponse or len(reponse.strip()) == 0:
        return False, "Réponse vide"

    nb_mots = len(reponse.split())

    # Ratio mots/chars : si < 2%, c'est quasi que du formatage
    if len(reponse) > 1000:
        ratio = nb_mots / len(reponse)
        if ratio < 0.02:
            issues.append(
                f"Ratio mots/chars trop faible ({ratio:.4f} — probable boucle de formatage)"
            )

    # Détection de séparateurs de tableau excessifs
    nb_separateurs = reponse.count("---")
    if nb_separateurs > 50:
        issues.append(f"Trop de séparateurs de tableau ({nb_separateurs})")

    # Détection de lignes dupliquées
    lignes = [l.strip() for l in reponse.split("\n") if l.strip()]
    if len(lignes) > 10:
        unique = set(lignes)
        if len(unique) < len(lignes) * 0.3:
            issues.append(
                f"Plus de 70% de lignes dupliquées ({len(unique)} uniques / {len(lignes)} total)"
            )

    # v7.1 : Seuils minimum par chapitre (en mots)
    # Les chapitres multi-sections doivent produire plus de contenu
    SEUILS_MIN_MOTS = {
        "Chapitre 1 : Identité du Projet": 80,
        "Chapitre 2 : Architecture générale": 100,
        "Chapitre 3 : Description étape par étape": 250,  # 5 sections obligatoires
        "Chapitre 4 : Cartographie des Données (Technique)": 150,
        "Chapitre 5 : Recommandations": 100,
        "Chapitre 6 : Maintenabilité": 150,
    }
    seuil = SEUILS_MIN_MOTS.get(titre, 50)
    if nb_mots < seuil:
        issues.append(
            f"Trop court pour ce chapitre ({nb_mots} mots, minimum attendu : {seuil})"
        )

    # v7.1 : Vérification des sections manquantes pour le chapitre 3
    if titre == "Chapitre 3 : Description étape par étape":
        sections_attendues = ["vigilance", "interface", "inter-outil"]
        texte_lower = reponse.lower()
        sections_trouvees = sum(1 for s in sections_attendues if s in texte_lower)
        if sections_trouvees == 0:
            issues.append(
                "Chapitre 3 incomplet : il manque les sections Points de vigilance / Interfaces Inter-outils"
            )

    if issues:
        return False, "; ".join(issues)
    return True, "OK"


# ==========================================================
# INTERROGER L'AGENT SSE (avec retry) — v7 amélioré
# ==========================================================
def interroger_agent_sse(
    question, titre_chapitre="", tentative_num=1, max_tentatives=3
):
    log_separateur(
        f"REQUÊTE SSE — {titre_chapitre or 'sans titre'} (tentative {tentative_num}/{max_tentatives})"
    )
    log(f"Prompt envoyé ({len(question)} chars) :")
    log(f"  {question[:800]}")
    if len(question) > 800:
        log(f"  ... [tronqué, longueur totale : {len(question)}]")

    body = {
        "text": question,
        "projectId": PROJECT_ID,
        "userId": RUN_ID,
        "sse": True,
        "projectConfigOverride": {"ai": {"max_tokens": 4000, "history": False}},
    }

    log_json("Body de la requête", body)

    t0 = time.time()
    try:
        response = requests.post(
            URL_QUERY, headers=HEADERS, json=body, stream=True, timeout=300
        )
        log(f"→ HTTP {response.status_code}")
    except Exception as e:
        log(f"❌ Erreur connexion SSE : {e}", "ERROR")
        return ""

    if response.status_code != 200:
        log(
            f"❌ Statut inattendu : {response.status_code} — body : {response.text[:500]}",
            "ERROR",
        )
        return ""

    client = sseclient.SSEClient(response)
    reponses_par_message = {}
    reponse_complete = ""
    nb_events = 0
    nb_chunks = 0
    sources_finales = []
    search_results_finaux = []
    dernier_message_id = None
    raw_events_log = []
    nb_passes = 0

    for event in client.events():
        nb_events += 1
        if event.data:
            try:
                data_json = json.loads(event.data)

                if nb_events <= 5:
                    raw_events_log.append({"event_n": nb_events, "data": data_json})

                event_type = data_json.get("type", "")

                if event_type == "callLLMWithTools":
                    msg_id = data_json.get("messageId")
                    if msg_id and msg_id not in reponses_par_message:
                        nb_passes += 1
                        reponses_par_message[msg_id] = ""
                        log(
                            f"   🔄 Nouveau cycle LLM (passe {nb_passes}) — messageId: {msg_id}"
                        )

                if event_type == "empty-answer":
                    finish = data_json.get("finishReasons", [])
                    log(f"   ℹ️  empty-answer — finishReasons: {finish}")

                # Log des résultats d'outils RAG
                if "activity" in data_json:
                    for act in data_json.get("activity", []):
                        if act.get("type") == "toolResult":
                            raw_content = act.get("raw", {}).get("content", "")
                            if "No result" in str(raw_content):
                                log(f"   ⚠️  TOOL RESULT : 'No result'", "WARN")
                            else:
                                content_preview = str(raw_content)[:200]
                                log(f"   📄 TOOL RESULT : {content_preview}")

                if "answer" in data_json:
                    msg_id = data_json.get("messageId", "unknown")
                    nouvelle_valeur = data_json["answer"]
                    ancienne_valeur = reponses_par_message.get(msg_id, "")
                    if nouvelle_valeur != ancienne_valeur:
                        nb_chunks += 1
                        reponses_par_message[msg_id] = nouvelle_valeur

                        # v7 : seuil anti-boucle par taille
                        if len(nouvelle_valeur) > 12000:
                            log(
                                f"   ⚠️  Réponse > 12000 chars ({len(nouvelle_valeur)}) — interruption préventive",
                                "WARN",
                            )
                            break

                        # v7 : détection boucle par ratio mots/chars en temps réel
                        if len(nouvelle_valeur) > 5000:
                            nb_mots_courant = len(nouvelle_valeur.split())
                            if nb_mots_courant < 100:
                                log(
                                    f"   ⚠️  Boucle de formatage détectée ({nb_mots_courant} mots / {len(nouvelle_valeur)} chars) — interruption",
                                    "WARN",
                                )
                                break

                if data_json.get("end") is True and "messageId" in data_json:
                    dernier_message_id = data_json.get("messageId")
                    sources_obj = data_json.get("sources", {})
                    if sources_obj:
                        sources_finales = sources_obj.get("values", [])
                    search_obj = data_json.get("search", {})
                    if search_obj:
                        search_results_finaux = search_obj.get("results", [])
                    raw_events_log.append(
                        {"event_n": nb_events, "data_final": data_json}
                    )

            except Exception as e:
                if nb_events <= 5:
                    raw_events_log.append(
                        {"event_n": nb_events, "raw": str(event.data)[:200]}
                    )
                continue

    if reponses_par_message:
        reponse_complete = max(reponses_par_message.values(), key=len)
        log(
            f"   📊 Passes LLM : {nb_passes} | Réponses : { {k: len(v) for k, v in reponses_par_message.items()} }"
        )
        log(f"   ✅ Réponse retenue : {len(reponse_complete)} chars")

    duree_s = time.time() - t0

    reponse_complete = nettoyer_reponse_boucle(reponse_complete)

    log("── PREMIERS EVENTS SSE BRUTS ────────────────────────────────────────")
    for ev in raw_events_log[:6]:
        log(f"  {json.dumps(ev, ensure_ascii=False)[:300]}")

    log(f"messageId final : {dernier_message_id}")

    log_reponse_sse(
        reponse_complete=reponse_complete,
        sources=sources_finales,
        search_results=search_results_finaux,
        nb_events=nb_events,
        nb_chunks=nb_chunks,
        duree_s=duree_s,
        titre_chapitre=titre_chapitre,
    )

    # Retry si réponse insuffisante
    nb_placeholders = reponse_complete.lower().count("donnée non identifiée")
    reponse_trop_courte = len(reponse_complete) < 300
    reponse_que_placeholders = nb_placeholders > 5

    if (
        reponse_trop_courte or reponse_que_placeholders
    ) and tentative_num < max_tentatives:
        log(
            f"   🔄 Réponse insuffisante. Retry dans 30s... (tentative {tentative_num + 1}/{max_tentatives})",
            "WARN",
        )
        time.sleep(30)
        return interroger_agent_sse(
            question,
            titre_chapitre=titre_chapitre,
            tentative_num=tentative_num + 1,
            max_tentatives=max_tentatives,
        )

    return reponse_complete


# ==========================================================
# SUPPRIMER DOCUMENTS DU RAG
# ==========================================================
def supprimer_documents(ids):
    log_separateur("SUPPRESSION DES DOCUMENTS DU RAG")
    for doc_id in ids:
        url = f"{URL_DELETE}?projectId={PROJECT_ID}&id={doc_id}"
        response = requete_avec_retry("DELETE", url, headers=HEADERS)
        if response:
            log(f"🗑  Supprimé : {doc_id}")
        else:
            log(f"⚠️  Échec suppression : {doc_id}", "WARN")
    log("Suppression terminée")


# ==========================================================
# PROMPTS ENRICHIS v7 — Simplifiés + anti-boucle + XML structuré
# ==========================================================

# Instruction anti-boucle commune à tous les chapitres
REGLE_ANTI_BOUCLE = (
    "\n<regle_formatage>\n"
    "RÈGLES MARKDOWN OBLIGATOIRES :\n"
    "- Séparateurs de tableau : EXACTEMENT 3 tirets par colonne (|---|). Jamais plus.\n"
    "- Maximum 20 lignes par tableau. Si plus de données, résume le reste en texte.\n"
    "- Maximum 6 colonnes par tableau.\n"
    "- Cellules courtes : pas de phrases longues dans les cellules de tableau.\n"
    "- Si tu détectes que tu répètes le même pattern, STOP immédiatement.\n"
    "</regle_formatage>\n"
)


def construire_prompt_chapitre(
    titre_chapitre, noms_documents_str, num_dossier, contexte_cumule=""
):
    """
    v7.1 : Réintroduction d'une incitation RAG légère pour les chapitres lourds.
    Le prompt système contient l'instruction, mais Gemini Flash la court-circuite
    quand le contexte cumulé lui semble suffisant. On ajoute un rappel ciblé.
    """

    # Chapitres nécessitant impérativement une consultation RAG approfondie
    CHAPITRES_RAG_INTENSIF = {
        "Chapitre 3 : Description étape par étape",
        "Chapitre 4 : Cartographie des Données (Technique)",
        "Chapitre 6 : Maintenabilité",
    }

    # Contexte de base
    contexte_base = (
        f"<contexte>\n"
        f"Projet dossier {num_dossier}.\n"
        f"Fichiers disponibles dans le RAG : {noms_documents_str}.\n"
    )

    # Ajout d'une incitation RAG ciblée pour les chapitres lourds
    if titre_chapitre in CHAPITRES_RAG_INTENSIF:
        contexte_base += (
            f"IMPORTANT : Consulte les documents du RAG via l'outil de recherche "
            f"avant de rédiger. Ce chapitre nécessite des données factuelles "
            f"extraites des fichiers sources.\n"
        )

    contexte_base += f"</contexte>\n"

    if contexte_cumule:
        contexte_base += (
            f"\n<chapitres_precedents>\n{contexte_cumule}\n</chapitres_precedents>\n"
        )

    # v7 : Structures spécifiques par chapitre — simplifiées, avec XML
    structures = {
        "Chapitre 1 : Identité du Projet": (
            "<tache>Génère le contenu du Chapitre 1 : Identité du Projet.</tache>\n"
            "<structure>\n"
            "Remplis ces sections :\n"
            "- **Nom du projet** : Extrais le nom exact depuis ficheStarcommand.txt\n"
            "- **Auteur(s)** : Propriétaire et Co-propriétaire depuis ficheStarcommand.txt\n"
            "- **Mission** : Résumé détaillé du besoin métier en 3-5 phrases. "
            "Explique POURQUOI ce projet existe, quel problème il résout, pour qui.\n"
            "- **Périmètre** : Fréquence d'usage, population cible, type de données traitées "
            "(nature, sensibilité RGPD, volume).\n"
            "</structure>\n"
            "Commence directement par le contenu. Pas de titre de chapitre."
        ),
        "Chapitre 2 : Architecture générale": (
            "<tache>Génère le contenu du Chapitre 2 : Architecture générale.</tache>\n"
            "<structure>\n"
            "- **Technologies** : Pour CHAQUE technologie, donne le nom, "
            "son rôle précis dans le projet (1-2 phrases).\n"
            "- **Schéma de flux (Mermaid)** : Diagramme graph LR avec subgraph "
            "Chaque arrête (edge) doit être numéroté dans l'ordre "
            'du flux dans son libellé, exemple : A -- "1. Dépose fichier" --> G.'
            "Nœuds avec guillemets doubles. 8-12 étapes numérotées.\n"
            "</structure>\n"
            "Commence directement par le contenu. Pas de titre de chapitre."
        ),
        "Chapitre 3 : Description étape par étape": (
            "<tache>Génère le contenu COMPLET du Chapitre 3...\n"
            "INTERDIT : Ne déduis pas les étapes depuis le schéma Mermaid du chapitre 2. "
            "Tu DOIS interroger le RAG pour lire les fichiers."
            "Mode opératoire avant de rédiger. Le workflow doit comporter 15 à 25 étapes "
            "granulaires reflétant la logique complète du projet (boucles, conditions, appels "
            "de macros, gestion des fichiers). "
            "Ce chapitre comporte 4 sections obligatoires (ou 5 si UiPath détecté). "
            "Ne t'arrête pas après le tableau Workflow.</tache>\n"
            "<structure>\n"
            "SECTION 1 — **Workflow** : Tableau avec colonnes : Étape | Outil | Action réalisée | Ce qui est impacté.\n"
            "Prend tout le RAG en considération pour rédiger un worklow de maximum 25 lignes. Cellules courtes et précises.\n"
            "\n"
            "SECTION 2 — **Points de vigilance** : Boucles ForEach, conditions If/Case, fragilités. "
            "Développe chaque point en 1-2 phrases de texte (pas de tableau).\n"
            "\n"
            "SECTION 3 — **Interfaces Inter-outils** : Pour chaque transfert de données entre outils, "
            "indique le mécanisme (fichier sur réseau, ouverture Excel, copier-coller, API...) "
            "et si c'est Synchrone ou Asynchrone. Rédige en texte ou en liste à puces.\n"
            "\n"
            "SECTION 4 — **Spécificités Macro Excel** (si VBA détecté dans les fichiers .xlsm) : "
            "Tableau avec Feuille | Macro/Formule | Description de la logique. "
            "Consulte les fichiers _parsed.json dans le RAG pour extraire le code VBA.\n"
            "\n"
            "SECTION 5 — **Spécificités UiPath** (si Main.xaml détecté) : Séquences principales, "
            "variables clés, interactions avec Excel. Consulte Main.xaml_parsed.json dans le RAG.\n"
            "</structure>\n"
            "Commence directement par le contenu. Pas de titre de chapitre. "
            "Tu DOIS remplir toutes les sections applicables."
        ),
        "Chapitre 4 : Cartographie des Données (Technique)": (
            "<tache>Génère le contenu du Chapitre 4 : Cartographie des Données.</tache>\n"
            "<structure>\n"
            "- **Sources & Destinations** : Pour chaque source/destination :\n"
            "  Nom | Type (API, Fichier, Réseau, Teams...) | Mode (Lecture/Écriture) | Description.\n"
            "- **Dictionnaire des données et Variables** : Tableau avec :\n"
            "  Fichier | Nom de la variable | Type (String, Int, DataTable...) | Utilisation.\n"
            "  Extrais depuis Main.xaml_parsed.json ET les fichiers Excel.\n"
            "</structure>\n"
            "Commence directement par le contenu. Pas de titre de chapitre."
        ),
        "Chapitre 5 : Recommandations": (
            "<tache>Génère le contenu du Chapitre 5 : Recommandations.</tache>\n"
            "<structure>\n"
            "Produis un tableau avec exactement ces colonnes :\n"
            "| Domaine | Recommandation | Impact attendu | Priorité |\n"
            "| --- | --- | --- | --- |\n"
            "\n"
            "Produis entre 6 et 8 lignes couvrant : Performance, Robustesse, "
            "Maintenabilité, Sécurité/RGPD, Évolutivité.\n"
            "</structure>\n"
            "Commence directement par le tableau. Pas de titre de chapitre."
        ),
        "Chapitre 6 : Maintenabilité": (
            "<tache>Génère le contenu du Chapitre 6 : Maintenabilité.</tache>\n"
            "<structure>\n"
            "Sois concis. Chaque section en 2-4 phrases.\n"
            "- **Score de Complexité** : [Bas/Moyen/Haut] avec justification.\n"
            "- **Audit nécessaire ?** : [Oui/Non] avec justification.\n"
            "- **Indice de Maintenabilité** : Score XX/100 avec ce tableau :\n"
            "| Pilier | Score /25 | Justification |\n"
            "| --- | --- | --- |\n"
            "| Homogénéité | X/25 | ... |\n"
            "| Standardisation | X/25 | ... |\n"
            "| Documentation Source | X/25 | ... |\n"
            "| Dépendances Critiques | X/25 | ... |\n"
            "- **Points forts** : 3-4 points qui facilitent la reprise.\n"
            "- **Points de fragilité** : 3-4 points de risque technique.\n"
            "</structure>\n"
            "Commence directement par le contenu. Pas de titre de chapitre."
        ),
    }

    structure_chapitre = structures.get(
        titre_chapitre, f"<tache>Génère le contenu de '{titre_chapitre}'.</tache>"
    )

    # v7 : assemblage propre = contexte + structure + anti-boucle
    prompt_complet = f"{contexte_base}\n{structure_chapitre}\n{REGLE_ANTI_BOUCLE}"

    return prompt_complet


def construire_resume_chapitre(titre, reponse):
    """
    v7 : Résumé FILTRÉ — exclut les lignes de formatage (tirets, Mermaid, séparateurs)
    pour éviter de propager du bruit dans le contexte cumulatif.
    """
    lignes_utiles = []
    for ligne in reponse.split("\n"):
        ligne = ligne.strip()
        if not ligne:
            continue
        # Ignorer les séparateurs de tableau
        if ligne.startswith("|") and "---" in ligne:
            continue
        # Ignorer le code Mermaid
        if ligne.startswith("```"):
            continue
        if ligne.startswith("graph ") or ligne.startswith("subgraph"):
            continue
        # Ignorer les lignes quasi-vides (que des tirets/pipes)
        contenu = (
            ligne.replace("-", "").replace("|", "").replace(":", "").replace(" ", "")
        )
        if len(contenu) < 5 and len(ligne) > 10:
            continue
        lignes_utiles.append(ligne)
        if len(" ".join(lignes_utiles)) > 300:
            break

    resume = " ".join(lignes_utiles)[:300]
    if len(reponse) > 300:
        resume += "..."
    return f'<chapitre titre="{titre}">{resume}</chapitre>'


# ==========================================================
# SCRIPT PRINCIPAL
# ==========================================================
log_separateur("DÉMARRAGE DU SCRIPT PRINCIPAL")

# 1. Upload avec extraction docx enrichie
noms_documents, ids_documents_rag = uploader_les_documents(DOSSIER_DOCUMENTS)
log(f"Noms uploadés       : {noms_documents}")
log(f"IDs RAG suppression : {ids_documents_rag}")

noms_documents_str = ", ".join(noms_documents)

CHAPITRES_TITRES = [
    "Chapitre 1 : Identité du Projet",
    "Chapitre 2 : Architecture générale",
    "Chapitre 3 : Description étape par étape",
    "Chapitre 4 : Cartographie des Données (Technique)",
    "Chapitre 5 : Recommandations",
    "Chapitre 6 : Maintenabilité",
]

# v7 : prompt initial simplifié — le prompt système gère déjà le rôle
PROMPT_INITIAL = (
    f"J'ai déposé {len(noms_documents)} fichiers dans le RAG pour le dossier {NUM_DOSSIER} : "
    f"{noms_documents_str}. "
    f"Confirme que tu as accès aux documents en listant ceux que tu retrouves "
    f"avec un résumé de 1 phrase du contenu de chaque fichier."
)

# Init document Word
doc_word = Document()
doc_word.add_heading("Documentation Technique - StarDoc✨", level=0)
phrase_ia = doc_word.add_paragraph()
run = phrase_ia.add_run(
    "Brouillon de documentation générée par l'IA sur la base des documents fournis dans la fiche StarCommand. "
    "Pour remplacer les diagrammes coller le code Mermaid dans https://app.diagrams.net/ > Organiser > Insérer > Mermaid"
)
run.font.size = Pt(9)
run.italic = True
phrase_ia.alignment = 0

# Pause vectorisation (l'indexation est fiable côté Prisme, pas de vérification nécessaire)
log_separateur("PAUSE VECTORISATION RAG")
log(f"Attente {PAUSE_VECTORISATION}s pour indexation des documents...")
time.sleep(PAUSE_VECTORISATION)
log("Fin de la pause de vectorisation.")

# Envoi du contexte initial
log_separateur("ENVOI DU PROMPT INITIAL DE CONTEXTE")
reponse_contexte = interroger_agent_sse(
    PROMPT_INITIAL, titre_chapitre="CONTEXTE INITIAL"
)

if (
    not reponse_contexte
    or "InternalServerError" in reponse_contexte
    or len(reponse_contexte) < 5
):
    msg = "ERREUR FATALE : Le contexte initial a echoue. Arret du script."
    log(msg)
    print(msg)
    supprimer_documents(ids_documents_rag)
    sys.exit(1)

log(f"Contexte initial confirme ({len(reponse_contexte)} chars).")
time.sleep(5)

# Génération des chapitres avec contexte cumulatif et validation qualité
log_separateur("ÉTAPE 2 : GÉNÉRATION DES CHAPITRES")
contexte_cumule = ""

for i, titre_chapitre in enumerate(CHAPITRES_TITRES, start=1):
    log(f"\n{'=' * 60}")
    log(f"CHAPITRE {i}/{len(CHAPITRES_TITRES)} : {titre_chapitre}")
    log(f"{'=' * 60}")

    reponse = None
    max_quality_retries = 3

    for quality_attempt in range(1, max_quality_retries + 1):
        prompt = construire_prompt_chapitre(
            titre_chapitre, noms_documents_str, NUM_DOSSIER, contexte_cumule
        )

        # v7 : si retry qualité, ajouter instruction corrective
        if quality_attempt > 1:
            prompt += (
                "\n\n<correction>\n"
                "ATTENTION : Ta réponse précédente contenait trop de formatage "
                "et pas assez de contenu textuel. "
                "Privilégie le texte en prose et les tableaux COURTS (max 15 lignes). "
                "Utilise des séparateurs de tableau simples : |---|.\n"
                "</correction>"
            )

        reponse = interroger_agent_sse(prompt, titre_chapitre=titre_chapitre)

        # v7 : validation qualité
        valide, raison = valider_qualite_chapitre(titre_chapitre, reponse)
        if valide:
            log(
                f"   ✅ Qualité validée (tentative {quality_attempt}/{max_quality_retries})"
            )
            break
        else:
            log(
                f"   ⚠️  Qualité insuffisante (tentative {quality_attempt}/{max_quality_retries}): {raison}",
                "WARN",
            )
            if quality_attempt < max_quality_retries:
                log(f"   🔄 Retry qualité dans 15s...", "WARN")
                time.sleep(15)

    print("\n" + "=" * 80)
    print(f"  {titre_chapitre}")
    print("=" * 80)
    print(reponse)
    print("=" * 80 + "\n")

    doc_word.add_heading(titre_chapitre, level=1)
    ajouter_texte_markdown(doc_word, reponse)
    doc_word.add_page_break()
    log(f"📘 {titre_chapitre} intégré au document Word.")

    # v7 : résumé cumulatif filtré
    resume = construire_resume_chapitre(titre_chapitre, reponse)
    contexte_cumule += f"\n{resume}"

    # Pause entre chapitres
    if i < len(CHAPITRES_TITRES):
        time.sleep(3)

# Sauvegarde
log_separateur("ÉTAPE 3 : SAUVEGARDE ET NETTOYAGE")
chemin_sortie = os.path.join(DOSSIER_DOCUMENTS, NOM_FICHIER_SORTIE)
doc_word.save(chemin_sortie)
log(f"✅ Document Word sauvegardé : {chemin_sortie}")

# Nettoyage RAG
supprimer_documents(ids_documents_rag)

# Nettoyage JSON parsés
for f in os.listdir(DOSSIER_DOCUMENTS):
    if f.endswith("_parsed.json"):
        try:
            os.remove(os.path.join(DOSSIER_DOCUMENTS, f))
            log(f"🗑  JSON supprimé : {f}")
        except Exception as e:
            log(f"⚠️  Impossible de supprimer {f} : {e}", "WARN")

log_separateur("FIN DU SCRIPT")
log(f"🎉 Script terminé. Log disponible : {_log_path}")

if _logger_file:
    _logger_file.close()
